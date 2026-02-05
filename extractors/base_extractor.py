"""
提取器基类
==========
公共方法、日志、异常处理
"""

import os
import re
from abc import ABC, abstractmethod
from typing import List, Dict, Optional, Any
from dataclasses import dataclass, field
from docx import Document

from .text_utils import parse_number, parse_floor, normalize_label, match_field, extract_year, extract_date
from .table_utils import (
    row_to_text_list, find_header_row, find_column_indices,
    find_row_by_label, find_rows_by_labels, auto_detect_tables,
    get_effective_row_cells, extract_property_rights_generic, get_row_cells_simple
)


@dataclass
class Position:
    """位置信息"""
    table_index: int = -1
    row_index: int = -1
    col_index: int = -1


@dataclass
class LocatedValue:
    """带位置的值"""
    value: Any = None
    position: Position = field(default_factory=Position)
    raw_text: str = ""


@dataclass
class Factor:
    """因素数据"""
    name: str = ""
    description: str = ""
    level: str = ""
    index: int = 100
    desc_pos: Position = field(default_factory=Position)
    level_pos: Position = field(default_factory=Position)
    index_pos: Position = field(default_factory=Position)


@dataclass
class ExtractionWarning:
    """提取警告"""
    field: str
    message: str
    table_idx: int = -1
    row_idx: int = -1
    col_idx: int = -1
    raw_text: str = ""


@dataclass
class ExtractionStats:
    """提取统计"""
    total_fields: int = 0
    extracted_fields: int = 0
    failed_fields: int = 0
    warnings: List[ExtractionWarning] = field(default_factory=list)


class BaseExtractor(ABC):
    """提取器基类"""

    # 子类需要定义的表格类型列表
    TABLE_TYPES: List[str] = []

    # 区域关键词映射
    DISTRICT_KEYWORDS = {
        '武进': '武进区',
        '新北': '新北区',
        '天宁': '天宁区',
        '钟楼': '钟楼区',
        '金坛': '金坛区',
        '溧阳': '溧阳市',
        '经开': '经开区',
    }

    # 因素名称分类
    LOCATION_FACTORS = ['区域位置', '楼幢位置', '朝向', '交通条件', '配套设施',
                        '环境质量', '景观', '物业管理', '繁华程度', '驻车条件']
    PHYSICAL_FACTORS = ['地形地势', '地质土壤', '开发程度', '建筑面积', '空间布局',
                        '新旧程度', '装饰装修', '建筑结构', '物业类型', '设施设备',
                        '套内建筑面积', '楼宇等级']
    RIGHTS_FACTORS = ['规划条件', '土地使用期限', '担保物权设立', '租赁占用状况',
                      '拖欠税费状况', '其他权益状况', '土地剩余使用年限', '登记状况',
                      '他项权利', '限制权利', '其他因素']

    def __init__(self, auto_detect: bool = True):
        self.doc = None
        self.tables = []
        self.full_text = ""
        self.auto_detect = auto_detect

        # 表格索引映射
        self.table_indices: Dict[str, int] = {}

        # 提取统计
        self.stats = ExtractionStats()

    def load_document(self, doc_path: str):
        """加载文档"""
        self.doc = Document(doc_path)
        self.tables = self.doc.tables
        self.full_text = "\n".join([p.text for p in self.doc.paragraphs])

        # 自动检测表格
        if self.auto_detect and self.TABLE_TYPES:
            self.table_indices = auto_detect_tables(self.tables, self.TABLE_TYPES)

    def get_table(self, table_type: str):
        """获取指定类型的表格"""
        idx = self.table_indices.get(table_type, -1)
        if 0 <= idx < len(self.tables):
            return self.tables[idx]
        return None

    def get_table_index(self, table_type: str) -> int:
        """获取表格索引"""
        return self.table_indices.get(table_type, -1)

    # ========================================================================
    # 安全提取方法
    # ========================================================================

    def safe_extract_number(self,
                            text: str,
                            field_name: str,
                            position: Position = None,
                            is_percentage: bool = False,
                            default: float = None) -> Optional[float]:
        """
        安全提取数值

        失败时记录警告而不是抛出异常
        """
        self.stats.total_fields += 1

        result = parse_number(text, is_percentage=is_percentage, default=default)

        if result is not None:
            self.stats.extracted_fields += 1
        else:
            self.stats.failed_fields += 1
            if text and text.strip() and text.strip() not in ['/', '-', '—', '']:
                self.stats.warnings.append(ExtractionWarning(
                    field=field_name,
                    message=f"无法解析数值",
                    table_idx=position.table_index if position else -1,
                    row_idx=position.row_index if position else -1,
                    col_idx=position.col_index if position else -1,
                    raw_text=text[:50] if text else ""
                ))

        return result

    def safe_extract_text(self,
                          text: str,
                          field_name: str,
                          position: Position = None) -> str:
        """安全提取文本"""
        self.stats.total_fields += 1

        if text and text.strip():
            self.stats.extracted_fields += 1
            return text.strip()
        else:
            self.stats.failed_fields += 1
            return ""

    def create_located_value(self,
                             value: Any,
                             table_idx: int,
                             row_idx: int,
                             col_idx: int,
                             raw_text: str = "") -> LocatedValue:
        """创建带位置的值"""
        return LocatedValue(
            value=value,
            position=Position(table_idx, row_idx, col_idx),
            raw_text=raw_text or str(value) if value is not None else ""
        )

    # ========================================================================
    # 动态提取方法
    # ========================================================================

    def extract_row_values(self,
                           table,
                           table_idx: int,
                           row_idx: int,
                           col_indices: Dict[str, int],
                           is_numeric: bool = True,
                           is_percentage: bool = False) -> Dict[str, LocatedValue]:
        """
        从行中提取多列的值
        """
        result = {}

        if row_idx < 0 or row_idx >= len(table.rows):
            return result

        cells = row_to_text_list(table.rows[row_idx])

        for field_name, col_idx in col_indices.items():
            if col_idx < 0 or col_idx >= len(cells):
                continue

            raw_text = cells[col_idx]

            if is_numeric:
                value = self.safe_extract_number(
                    raw_text, field_name,
                    Position(table_idx, row_idx, col_idx),
                    is_percentage=is_percentage
                )
            else:
                value = self.safe_extract_text(
                    raw_text, field_name,
                    Position(table_idx, row_idx, col_idx)
                )

            if value is not None:
                result[field_name] = self.create_located_value(
                    value, table_idx, row_idx, col_idx, raw_text
                )

        return result

    def extract_column_values(self,
                              table,
                              table_idx: int,
                              col_idx: int,
                              row_mapping: Dict[str, int],
                              is_numeric: bool = True,
                              is_percentage: bool = False) -> Dict[str, LocatedValue]:
        """
        从列中提取多行的值
        """
        result = {}

        for field_name, row_idx in row_mapping.items():
            if row_idx < 0 or row_idx >= len(table.rows):
                continue

            cells = row_to_text_list(table.rows[row_idx])
            if col_idx < 0 or col_idx >= len(cells):
                continue

            raw_text = cells[col_idx]

            if is_numeric:
                value = self.safe_extract_number(
                    raw_text, field_name,
                    Position(table_idx, row_idx, col_idx),
                    is_percentage=is_percentage
                )
            else:
                value = self.safe_extract_text(
                    raw_text, field_name,
                    Position(table_idx, row_idx, col_idx)
                )

            if value is not None:
                result[field_name] = self.create_located_value(
                    value, table_idx, row_idx, col_idx, raw_text
                )

        return result

    # ========================================================================
    # 全文正则提取
    # ========================================================================

    def extract_from_text(self, patterns: Dict[str, List[str]]) -> Dict[str, str]:
        """
        从全文中用正则提取

        Args:
            patterns: {field_name: [regex_pattern1, regex_pattern2, ...], ...}

        Returns:
            {field_name: extracted_value, ...}
        """
        result = {}

        for field_name, pattern_list in patterns.items():
            for pattern in pattern_list:
                try:
                    match = re.search(pattern, self.full_text)
                    if match:
                        result[field_name] = match.group(1).strip()
                        break
                except Exception:
                    continue

        return result

    # ========================================================================
    # 区域解析
    # ========================================================================

    def parse_district(self, address: str) -> Dict[str, str]:
        """
        从地址解析区域信息

        Returns:
            {'district': '武进区', 'street': '湖塘镇'}
        """
        result = {'district': '', 'street': ''}

        if not address:
            return result

        # 解析区域
        for keyword, district in self.DISTRICT_KEYWORDS.items():
            if keyword in address:
                result['district'] = district
                break

        # 解析街道/镇
        street_match = re.search(r'([^\s区市县]+(?:镇|街道|乡))', address)
        if street_match:
            result['street'] = street_match.group(1)

        return result

    # ========================================================================
    # 因素分类
    # ========================================================================

    def classify_factor(self, factor_name: str) -> str:
        """
        分类因素名称

        Returns:
            'location' | 'physical' | 'rights' | 'unknown'
        """
        name = normalize_label(factor_name)

        if any(f in name for f in self.LOCATION_FACTORS):
            return 'location'
        elif any(f in name for f in self.PHYSICAL_FACTORS):
            return 'physical'
        elif any(f in name for f in self.RIGHTS_FACTORS):
            return 'rights'
        else:
            return 'unknown'

    def normalize_factor_name(self, name: str) -> str:
        """标准化因素名称"""
        return normalize_label(name)

    # ========================================================================
    # 抽象方法（子类必须实现）
    # ========================================================================

    @abstractmethod
    def extract(self, doc_path: str):
        """提取报告"""
        pass

    # ========================================================================
    # 输出统计
    # ========================================================================

    def print_stats(self):
        """打印提取统计"""
        print(f"\n📊 提取统计:")
        print(f"   总字段: {self.stats.total_fields}")
        print(f"   成功: {self.stats.extracted_fields}")
        print(f"   失败: {self.stats.failed_fields}")

        if self.stats.warnings:
            print(f"\n⚠️ 警告 ({len(self.stats.warnings)}):")
            for w in self.stats.warnings[:10]:
                print(f"   - [{w.field}] {w.message} (原文: {w.raw_text})")
            if len(self.stats.warnings) > 10:
                print(f"   ... 还有 {len(self.stats.warnings) - 10} 条警告")

    def get_stats_dict(self) -> Dict[str, Any]:
        """获取统计信息字典"""
        return {
            'total_fields': self.stats.total_fields,
            'extracted_fields': self.stats.extracted_fields,
            'failed_fields': self.stats.failed_fields,
            'success_rate': self.stats.extracted_fields / max(self.stats.total_fields, 1),
            'warning_count': len(self.stats.warnings),
        }
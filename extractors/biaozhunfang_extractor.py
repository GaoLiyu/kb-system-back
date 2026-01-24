"""
标准房报告精确提取器
====================
针对标准房报告的表格结构精确提取
表格索引（基于分析结果）：
- 表格6: 主要信息表（地址、面积、各类修正系数等，4个可比实例）
- 表格19: 详细因素比较表
- 表格20: 修正计算表
"""

import os
import re
from typing import Dict, List
from docx import Document
from dataclasses import dataclass, field


@dataclass
class Position:
    """位置信息"""
    table_index: int = -1
    row_index: int = -1
    col_index: int = -1


@dataclass
class LocatedValue:
    """带位置的值"""
    value: any = None
    position: Position = field(default_factory=Position)
    raw_text: str = ""


@dataclass
class Case:
    """可比实例"""
    case_id: str = ""  # A/B/C/D
    address: LocatedValue = field(default_factory=LocatedValue)
    data_source: str = ""
    building_area: LocatedValue = field(default_factory=LocatedValue)
    transaction_price: LocatedValue = field(default_factory=LocatedValue)  # 交易单价

    # 标准房特有的修正系数
    structure_factor: LocatedValue = field(default_factory=LocatedValue)  # 结构修正
    floor_factor: LocatedValue = field(default_factory=LocatedValue)  # 层次修正
    orientation_factor: LocatedValue = field(default_factory=LocatedValue)  # 朝向修正
    age_factor: LocatedValue = field(default_factory=LocatedValue)  # 成新修正
    east_to_west_factor: LocatedValue = field(default_factory=LocatedValue)  # 东西至修正
    physical_composite: LocatedValue = field(default_factory=LocatedValue)  # 实体状况综合

    # 计算表中的修正
    p1_transaction: str = ""  # P1交易情况修正
    p2_date: str = ""  # P2交易日期修正
    p3_physical: str = ""  # P3实体因素修正
    p4_location: str = ""  # P4区位状况修正
    composite_result: LocatedValue = field(default_factory=LocatedValue)  # P1×P2×P3×P4结果
    vs_result: LocatedValue = field(default_factory=LocatedValue)  # Vs×结果
    decoration_price: LocatedValue = field(default_factory=LocatedValue)  # 装修重置价
    attachment_price: LocatedValue = field(default_factory=LocatedValue)  # 附属物单价
    final_price: LocatedValue = field(default_factory=LocatedValue)  # 比准价格

    cart_type: str = ""  # 证号类型
    cart_code: str = ""  # 证号编码
    district: str = ""  # 区域（区/县）
    street: str = ""  # 街道/镇
    build_year: int = 0  # 建成年份
    total_floor: int = 0  # 总楼层
    current_floor: int = 0  # 所在楼层
    orientation: str = ""  # 朝向（文本描述）
    decoration: str = ""  # 装修状况
    structure: str = ""  # 建筑结构
    usage: str = ""  # 房屋性质
    transaction_date: str = ""  # 交易日期
    location_code: str = ""  # 区号
    east_to_west: str = ""  # 东西至
    appendages: str = ""  # 附属物
    avg_listing_price: LocatedValue = field(default_factory=LocatedValue) # 对应时点本片区二手房挂牌均价


@dataclass
class Subject:
    """估价对象（标准房 - 增强版）"""
    address: LocatedValue = field(default_factory=LocatedValue)
    building_area: LocatedValue = field(default_factory=LocatedValue)
    unit_price: LocatedValue = field(default_factory=LocatedValue)  # 评估单价
    total_price: LocatedValue = field(default_factory=LocatedValue)  # 评估总价
    transaction_price: LocatedValue = field(default_factory=LocatedValue)  # 交易单价

    # 修正系数
    structure_factor: LocatedValue = field(default_factory=LocatedValue)
    floor_factor: LocatedValue = field(default_factory=LocatedValue)
    orientation_factor: LocatedValue = field(default_factory=LocatedValue)
    age_factor: LocatedValue = field(default_factory=LocatedValue)
    east_to_west_factor: LocatedValue = field(default_factory=LocatedValue)
    physical_composite: LocatedValue = field(default_factory=LocatedValue)

    cart_type: str = "" # 证号类型
    cart_code: str = "" # 证号编码
    district: str = ""  # 区域（区/县）
    street: str = ""  # 街道/镇
    build_year: int = 0  # 建成年份
    total_floor: int = 0  # 总楼层
    current_floor: int = 0  # 所在楼层
    orientation: str = ""  # 朝向（文本描述）
    decoration: str = ""  # 装修状况
    structure: str = ""  # 建筑结构
    usage: str = ""  # 房屋性质
    transaction_date: str = ""  # 价值时点 && 交易时间
    location_code: str = "" # 区号
    east_to_west: str = "" # 东西至
    appendages: str = "" # 附属物
    avg_listing_price: LocatedValue = field(default_factory=LocatedValue)  # 对应时点本片区二手房挂牌均价

    appraisal_purpose: str = ""  # 估价目的


@dataclass
class BiaozhunfangExtractionResult:
    """标准房报告提取结果（增强版）"""
    source_file: str = ""
    subject: Subject = field(default_factory=Subject)
    cases: List[Case] = field(default_factory=list)

    # 最终结果（比准价格的平均值或加权值）
    final_price: LocatedValue = field(default_factory=LocatedValue)

    type: str = ""  # 类型


class BiaozhunfangExtractor:
    """标准房报告提取器（增强版）"""

    # 表格索引（默认值，会被自动检测覆盖）
    TABLE_MAIN_INFO = 6  # 主要信息表（34行）
    TABLE_DETAIL = 19  # 详细因素表（30行）
    TABLE_CORRECTION = 20  # 修正计算表（11行）
    TABLE_RESULT_SUMMARY = 2  # 结果汇总表

    def __init__(self, auto_detect: bool = True):
        self.doc = None
        self.tables = []
        self.full_text = ""  # 新增：完整文本用于正则提取
        self.auto_detect = auto_detect

    def extract(self, doc_path: str) -> BiaozhunfangExtractionResult:
        """提取标准房报告"""
        self.doc = Document(doc_path)
        self.tables = self.doc.tables
        self.full_text = "\n".join([p.text for p in self.doc.paragraphs])

        result = BiaozhunfangExtractionResult(source_file=os.path.basename(doc_path), type="biaozhunfang")

        print(f"\n📊 提取标准房报告: {os.path.basename(doc_path)}")
        print(f"   表格数量: {len(self.tables)}")

        # 自动检测表格索引
        if self.auto_detect:
            self._auto_detect_table_indices()
            print(f"   ✓ 自动检测: 主表={self.TABLE_MAIN_INFO}, 详细={self.TABLE_DETAIL}, 修正={self.TABLE_CORRECTION}")

        # 初始化4个可比实例
        result.cases = [Case(case_id='A'), Case(case_id='B'),
                        Case(case_id='C'), Case(case_id='D')]

        # 1. 提取结果汇总表（单价、总价）
        self._extract_result_summary(result)
        print(f"   ✓ 结果汇总表")

        # 2. 提取主要信息表
        self._extract_basic_table(result)
        print(f"   ✓ 主要信息表: 地址、面积、单价、总价")

        # 2. 从详细因素表提取基本信息和修正系数
        self._extract_detail_table(result)
        print(f"   ✓ 修正系数")

        # 3. 从修正计算表提取修正计算
        self._extract_correction_table(result)
        print(f"   ✓ 修正计算表: 比准价格")

        # 4. 提取扩展信息（估价目的）
        self._extract_extended_info(result)
        print(f"   ✓ 扩展信息: 估价目的")

        # 5. 解析区域信息
        self._parse_district(result)
        print(f"   ✓ 区域解析: {result.subject.district} {result.subject.street}")

        return result

    def _auto_detect_table_indices(self):
        """自动检测关键表格的索引位置（打分制，更稳）"""

        def norm(s: str) -> str:
            if not s:
                return ""
            return (
                s.replace("\u3000", " ")
                .replace("\n", " ")
                .replace("\t", " ")
                .strip()
            )

        def table_block(table, max_rows=6, max_cols=10) -> str:
            parts = []
            rN = min(len(table.rows), max_rows)
            for r in range(rN):
                row = table.rows[r]
                cN = min(len(row.cells), max_cols)
                for c in range(cN):
                    parts.append(norm(row.cells[c].text))
            # 拼成一个大块，便于 contains 判断
            return " ".join([p for p in parts if p])

        def contains_all(text: str, keys) -> bool:
            return all(k in text for k in keys)

        def count_hits(text: str, keys) -> int:
            return sum(1 for k in keys if k in text)

        # 候选分数
        best = {
            "result": (-1, self.TABLE_RESULT_SUMMARY),
            "main": (-1, self.TABLE_MAIN_INFO),
            "detail": (-1, self.TABLE_DETAIL),
            "corr": (-1, self.TABLE_CORRECTION),
        }

        for i, table in enumerate(self.tables):
            if len(table.rows) == 0:
                continue

            rows = len(table.rows)
            cols = len(table.columns) if table.columns else 0

            block = table_block(table, max_rows=8, max_cols=12)

            # 统一做一次“无空格版”，应对“P1 交易情况 修正”这种被拆开的情况
            block_compact = block.replace(" ", "")

            # ---------- 1) 结果汇总表 ----------
            # 特征：短小（行数少），且同时出现“总价/单价”
            score_result = 0
            if rows <= 6 and cols >= 3:
                if ("总价" in block_compact) and ("单价" in block_compact):
                    score_result += 10
                score_result += count_hits(block_compact, ["评估结果", "元/㎡", "万元"])
            if score_result > best["result"][0]:
                best["result"] = (score_result, i)

            # ---------- 2) 修正计算表 ----------
            # 特征：P1/P2/P3/P4 + 比准价格 等关键词非常强
            score_corr = 0
            corr_keys_strong = ["P1交易情况修正", "P2交易日期修正", "P3", "P4", "比准价格"]
            corr_keys_weak = ["Vs", "装修重置价", "附属物", "P1×P2×P3×P4"]
            if rows >= 8 and cols >= 5:
                score_corr += 2 * count_hits(block_compact, corr_keys_strong)
                score_corr += 1 * count_hits(block_compact, corr_keys_weak)
            if score_corr > best["corr"][0]:
                best["corr"] = (score_corr, i)

            # ---------- 3) A-D 对比类表（主表/详细表都会中）----------
            has_ad_header = (
                    ("估价对象" in block_compact)
                    and ("可比实例A" in block_compact)
                    and ("可比实例B" in block_compact)
                    and ("可比实例C" in block_compact)
                    and ("可比实例D" in block_compact)
            )

            if has_ad_header:
                # ---------- 3.1) 详细因素表 ----------
                # 特征：大量“修正系数/区位因素/综合得分”等字段
                score_detail = 0
                detail_strong = ["结构修正系数", "层次修正系数", "朝向修正系数", "成新修正系数", "实体状况系数综合"]
                detail_weak = ["区位因素", "区位状况综合", "修正结果", "装修成新率", "装修档次"]
                score_detail += 2 * count_hits(block_compact, detail_strong)
                score_detail += 1 * count_hits(block_compact, detail_weak)
                # 一般详细表列数较稳定（你这份是6列），但不要写死，只做加分
                if cols <= 7:
                    score_detail += 2
                if score_detail > best["detail"][0]:
                    best["detail"] = (score_detail, i)

                # ---------- 3.2) 主要信息表（主表） ----------
                # 特征：案件基础信息字段：案例来源、证号类型/编码、地址、建筑面积、交易时间等
                score_main = 0
                main_strong = ["案例来源", "证号类型", "证号编码", "地址", "建筑面积", "交易时间"]
                main_weak = ["结构", "层次", "朝向", "建成时间", "装修", "房屋性质", "区位（代码）"]
                score_main += 2 * count_hits(block_compact, main_strong)
                score_main += 1 * count_hits(block_compact, main_weak)
                # 主表一般更大（你这份34行、9列），给个形态加分
                if rows >= 25:
                    score_main += 2
                if cols >= 8:
                    score_main += 2
                if score_main > best["main"][0]:
                    best["main"] = (score_main, i)

        # 落盘（如果某类没找到，保留默认值）
        if best["result"][0] >= 6:
            self.TABLE_RESULT_SUMMARY = best["result"][1]
        if best["corr"][0] >= 6:
            self.TABLE_CORRECTION = best["corr"][1]
        if best["detail"][0] >= 6:
            self.TABLE_DETAIL = best["detail"][1]
        if best["main"][0] >= 6:
            self.TABLE_MAIN_INFO = best["main"][1]

    def _extract_result_summary(self, result: BiaozhunfangExtractionResult):
        """提取结果汇总表（单价、总价）"""
        if len(self.tables) <= self.TABLE_RESULT_SUMMARY:
            return

        table = self.tables[self.TABLE_RESULT_SUMMARY]

        if len(table.rows) < 2:
            return

        # cell = [c.text.strip() for c in table.rows[1].cells]

        result.subject.unit_price = LocatedValue(
            value=float(table.rows[2].cells[2].text),
            position=Position(table_index=self.TABLE_RESULT_SUMMARY, row_index=2, col_index=2),
            raw_text=table.rows[1].cells[2].text.strip(),
        )

        result.subject.total_price = LocatedValue(
            value=float(table.rows[1].cells[2].text),
            position=Position(table_index=self.TABLE_RESULT_SUMMARY, row_index=1, col_index=2),
            raw_text=table.rows[1].cells[2].text.strip(),
        )

    def _extract_basic_table(self, result: BiaozhunfangExtractionResult):
        """提取主要信息表"""
        if len(self.tables) <= self.TABLE_MAIN_INFO:
            return

        table = self.tables[self.TABLE_MAIN_INFO]

        COL_SUBJECT = 3
        COL_A = 4
        COL_B = 5
        COL_C = 6
        COL_D = 7

        for row_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]

            if len(cells) < 6:
                continue

            label = (cells[0]).replace(' ', '').replace('\u3000', '')
            label2 = (cells[3]).replace(' ', '').replace('\u3000', '')

            # 案例来源
            if label2 == '估价对象' and '可比实例' not in cells[COL_A] and cells[COL_A] != '案例来源':
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.data_source = cells[col]

            elif '证' in label and '类型' in label:
                result.subject.cart_type = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.cart_type = cells[col]

            elif '证' in label and '编码' in label:
                result.subject.cart_code = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.cart_code = cells[col]

            elif '地址' in label or '坐落' in label:
                result.subject.address = LocatedValue(
                    value=cells[COL_SUBJECT],
                    position=Position(self.TABLE_MAIN_INFO, row_idx, COL_SUBJECT),
                    raw_text=cells[COL_SUBJECT]
                )
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.address = LocatedValue(
                            value=cells[col],
                            position=Position(self.TABLE_MAIN_INFO, row_idx, col),
                            raw_text=cells[col]
                        )

            elif '评估面积' in label or '建筑面积' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.building_area = LocatedValue(
                        value=cells[COL_SUBJECT],
                        position=Position(self.TABLE_MAIN_INFO, row_idx, COL_SUBJECT),
                        raw_text=cells[COL_SUBJECT]
                    )
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        try:
                            area = float(re.sub(r'[^\d.]', '', cells[col]))
                            case.building_area = LocatedValue(
                                value=area,
                                position=Position(self.TABLE_MAIN_INFO, row_idx, col),
                                raw_text=cells[col]
                            )
                        except:
                            pass

            elif '结构' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.structure = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.structure = cells[col]

            elif '层次' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.current_floor = cells[COL_SUBJECT].split('/')[0]
                    result.subject.total_floor = cells[COL_SUBJECT].split('/')[1]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.current_floor = cells[col].split('/')[0]
                        case.total_floor = cells[col].split('/')[1]

            elif '朝向' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.orientation = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.orientation = cells[col]

            elif '建成时间' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.build_year = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.build_year = cells[col]

            elif '东西' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.east_to_west = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.east_to_west = cells[col]

            elif '装修' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.decoration = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.decoration = cells[col]

            elif '附属物' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.appendages = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.appendages = cells[col]

            elif '区位' in label and '代码' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.location_code = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.location_code = cells[col]

            elif '房屋性质' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.usage = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.usage = cells[col]

            elif '交易时间' in label or '价值时点' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.transaction_date = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.transaction_date = cells[col]

            elif '交易单价' in label:
                if COL_SUBJECT < len(cells) and not result.subject.transaction_price:
                    result.subject.transaction_price = LocatedValue(
                        value=float(cells[COL_SUBJECT]),
                        position=Position(self.TABLE_MAIN_INFO, row_idx, COL_SUBJECT),
                        raw_text=cells[COL_SUBJECT]
                    )
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.transaction_price = LocatedValue(
                            value=float(cells[col]),
                            position=Position(self.TABLE_MAIN_INFO, row_idx, col),
                            raw_text=cells[col]
                        )

            elif '二手房' in label or '挂牌均价' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.avg_listing_price = LocatedValue(
                        value=float(cells[COL_SUBJECT]),
                        position=Position(self.TABLE_MAIN_INFO, row_idx, COL_SUBJECT),
                        raw_text=cells[COL_SUBJECT]
                    )
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.avg_listing_price = LocatedValue(
                            value=float(cells[col]),
                            position=Position(self.TABLE_MAIN_INFO, row_idx, col),
                            raw_text=cells[col]
                        )

    def _extract_detail_table(self, result: BiaozhunfangExtractionResult):
        """提取详细因素表"""
        if len(self.tables) <= self.TABLE_DETAIL:
            return

        table = self.tables[self.TABLE_DETAIL]

        COL_SUBJECT = 1
        COL_A = 2
        COL_B = 3
        COL_C = 4
        COL_D = 5

        for row_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]

            label = (cells[0]).replace(' ', '').replace('\u3000', '')

            if len(cells) < 5:
                continue

            elif '结构修正系数' in label:
                self._extract_factor_row(result, cells, row_idx, 'structure_factor', COL_SUBJECT, COL_A)

            elif '层次修正系数' in label:
                self._extract_factor_row(result, cells, row_idx, 'floor_factor', COL_SUBJECT, COL_A)

            elif '朝向修正系数' in label:
                self._extract_factor_row(result, cells, row_idx, 'orientation_factor', COL_SUBJECT, COL_A)

            elif '成新修正系数' in label:
                self._extract_factor_row(result, cells, row_idx, 'age_factor', COL_SUBJECT, COL_A)

            elif '东西至修正系数' in label:
                self._extract_factor_row(result, cells, row_idx, 'east_to_west', COL_SUBJECT, COL_A)

            elif '实体状况系数综合' in label:
                self._extract_factor_row(result, cells, row_idx, 'physical_composite', COL_SUBJECT, COL_A)

    def _extract_factor_row(self, result, cells, row_idx, factor_name, col_subject, col_a):
        """提取修正系数行"""
        if len(cells) > col_subject:
            try:
                value = float(cells[col_subject]) / 100 if float(cells[col_subject]) > 10 else float(cells[col_subject])
                setattr(result.subject, factor_name, LocatedValue(
                    value=value,
                    position=Position(self.TABLE_DETAIL, row_idx, col_subject),
                    raw_text=cells[col_subject]
                ))
            except:
                pass

        for i, case in enumerate(result.cases):
            col = col_a + i
            if col < len(cells):
                try:
                    value = float(cells[col])
                    setattr(case, factor_name, LocatedValue(
                        value=value,
                        position=Position(self.TABLE_DETAIL, row_idx, col),
                        raw_text=cells[col]
                    ))
                except:
                    pass

    def _extract_correction_table(self, result: BiaozhunfangExtractionResult):
        """提取修正计算表"""
        if len(self.tables) <= self.TABLE_CORRECTION:
            return

        table = self.tables[self.TABLE_CORRECTION]

        COL_A = 1
        COL_B = 2
        COL_C = 3
        COL_D = 4

        ROW_P1 = 2
        ROW_P2 = 3
        ROW_P3 = 4
        ROW_P4 = 5
        ROW_COMPOSITE = 6
        ROW_VS_RESULT = 7
        ROW_DECORATION = 8
        ROW_ATTACHMENT = 9
        ROW_FINAL = 10

        for row_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]

            if len(cells) < 5:
                continue

            elif row_idx == ROW_P1:
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.p1_transaction = cells[col]

            elif row_idx == ROW_P2:
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.p2_date = cells[col]

            elif row_idx == ROW_P3:
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.p3_physical = cells[col]

            elif row_idx == ROW_P4:
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.p4_location = cells[col]

            elif row_idx == ROW_COMPOSITE:
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        try:
                            case.composite_result = LocatedValue(
                                value=float(cells[col]),
                                position=Position(self.TABLE_CORRECTION, row_idx, col),
                                raw_text=cells[col]
                            )
                        except:
                            pass

            elif row_idx == ROW_VS_RESULT:
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        try:
                            case.vs_result = LocatedValue(
                                value=float(cells[col]),
                                position=Position(self.TABLE_CORRECTION, row_idx, col),
                                raw_text=cells[col]
                            )
                        except:
                            pass

            elif row_idx == ROW_DECORATION:
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        try:
                            case.decoration_price = LocatedValue(
                                value=float(cells[col]),
                                position=Position(self.TABLE_CORRECTION, row_idx, col),
                                raw_text=cells[col]
                            )
                        except:
                            pass

            elif row_idx == ROW_ATTACHMENT:
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        try:
                            case.attachment_price = LocatedValue(
                                value=float(cells[col]),
                                position=Position(self.TABLE_CORRECTION, row_idx, col),
                                raw_text=cells[col]
                            )
                        except:
                            pass

            elif row_idx == ROW_FINAL:
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        try:
                            case.final_price = LocatedValue(
                                value=float(cells[col]),
                                position=Position(self.TABLE_CORRECTION, row_idx, col),
                                raw_text=cells[col]
                            )
                        except:
                            pass

    def _extract_extended_info(self, result: BiaozhunfangExtractionResult):
        """提取扩展信息（估价目的等）"""
        # 估价目的
        purpose_patterns = [
            r'估价目的[：:是为]*(.{5,80}?)(?:。|$)',
            r'本次估价目的是(.{5,80}?)(?:。|$)',
        ]

        for pattern in purpose_patterns:
            match = re.search(pattern, self.full_text)
            if match:
                result.subject.appraisal_purpose = match.group(1).strip()
                break

    def _parse_district(self, result: BiaozhunfangExtractionResult):
        """从地址解析区域信息"""
        address = result.subject.address.value or ""

        # 区/县
        district_patterns = [
            r'([\u4e00-\u9fa5]{2,4}区)',
            r'([\u4e00-\u9fa5]{2,4}县)',
            r'([\u4e00-\u9fa5]{2,4}市)',
        ]

        for pattern in district_patterns:
            match = re.search(pattern, address)
            if match:
                result.subject.district = match.group(1)
                break

        # 街道/镇
        street_patterns = [
            r'([\u4e00-\u9fa5]{2,6}街道)',
            r'([\u4e00-\u9fa5]{2,4}镇)',
            r'([\u4e00-\u9fa5]{2,4}乡)',
        ]

        for pattern in street_patterns:
            match = re.search(pattern, address)
            if match:
                result.subject.street = match.group(1)
                break

        # 同样处理可比实例
        for case in result.cases:
            case_addr = case.address.value or ""

            for pattern in district_patterns:
                match = re.search(pattern, case_addr)
                if match:
                    case.district = match.group(1)
                    break

            for pattern in street_patterns:
                match = re.search(pattern, case_addr)
                if match:
                    case.street = match.group(1)
                    break

# ============================================================================
# 测试
# ============================================================================

if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
    else:
        doc_path = "./uploads/kb_032-丰乐尚都---小高层.docx"

    extractor = BiaozhunfangExtractor()
    result = extractor.extract(doc_path)

    print(result.cases[0])
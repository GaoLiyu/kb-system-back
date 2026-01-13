"""
文档内容提取器
==============
提取Word文档的原文内容（段落和表格），用于前端展示
"""

from typing import List, Dict, Any
from docx import Document
from dataclasses import dataclass, field


@dataclass
class ContentItem:
    """内容项"""
    index: int                    # 序号
    type: str                     # paragraph / table
    text: str = ""                # 段落文本
    rows: List[List[str]] = None  # 表格行数据
    has_issue: bool = False       # 是否有问题（前端高亮用）
    issue_ids: List[int] = None   # 关联的问题ID


@dataclass
class DocumentContent:
    """文档内容"""
    filename: str
    contents: List[ContentItem] = field(default_factory=list)
    paragraph_count: int = 0
    table_count: int = 0


def extract_document_content(doc_path: str) -> DocumentContent:
    """
    提取文档原文内容

    Args:
        doc_path: 文档路径

    Returns:
        DocumentContent
    """
    import os
    from utils import convert_doc_to_docx

    # 处理doc文件
    if doc_path.lower().endswith('.doc'):
        doc_path = convert_doc_to_docx(doc_path)

    doc = Document(doc_path)
    result = DocumentContent(filename=os.path.basename(doc_path))

    # 获取文档的body元素
    body = doc.element.body

    index = 0
    para_count = 0
    table_count = 0

    # 按顺序遍历body的子元素
    for element in body:
        tag = element.tag.split('}')[-1]  # 去掉命名空间

        if tag == 'p':
            # 段落
            text = element.text or ''
            # 获取段落的完整文本（包括runs）
            for child in element.iter():
                if child.tag.endswith('}t'):
                    if child.text:
                        text = (text or '')

            # 从paragraph对象获取完整文本
            if para_count < len(doc.paragraphs):
                text = doc.paragraphs[para_count].text
                para_count += 1

            # 跳过空段落
            if text.strip():
                result.contents.append(ContentItem(
                    index=index,
                    type='paragraph',
                    text=text.strip(),
                    issue_ids=[]
                ))
                index += 1

        elif tag == 'tbl':
            # 表格
            if table_count < len(doc.tables):
                table = doc.tables[table_count]
                rows = []

                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace('\n', ' ')
                        row_data.append(cell_text)
                    rows.append(row_data)

                result.contents.append(ContentItem(
                    index=index,
                    type='table',
                    rows=rows,
                    issue_ids=[]
                ))
                index += 1
                table_count += 1

    result.paragraph_count = para_count
    result.table_count = table_count

    return result


def content_to_dict(content: DocumentContent) -> Dict[str, Any]:
    """转换为字典（用于API返回）"""
    return {
        'filename': content.filename,
        'paragraph_count': content.paragraph_count,
        'table_count': content.table_count,
        'contents': [
            {
                'index': item.index,
                'type': item.type,
                'text': item.text if item.type == 'paragraph' else None,
                'rows': item.rows if item.type == 'table' else None,
                'has_issue': item.has_issue,
                'issue_ids': item.issue_ids or [],
            }
            for item in content.contents
        ]
    }


def get_paragraphs_text(content: DocumentContent) -> List[Dict[str, Any]]:
    """
    获取所有段落文本（用于LLM审查）

    Returns:
        [{'index': 0, 'text': '...'}, ...]
    """
    return [
        {'index': item.index, 'text': item.text}
        for item in content.contents
        if item.type == 'paragraph' and item.text.strip()
    ]


def mark_issues(content: DocumentContent, issues: List[Dict]) -> DocumentContent:
    """
    标记有问题的内容项

    Args:
        content: 文档内容
        issues: 问题列表，每个问题需包含 paragraph_index 字段

    Returns:
        标记后的文档内容
    """
    # 建立索引到内容项的映射
    index_map = {item.index: item for item in content.contents}

    for i, issue in enumerate(issues):
        para_idx = issue.get('paragraph_index')
        if para_idx is not None and para_idx in index_map:
            item = index_map[para_idx]
            item.has_issue = True
            if item.issue_ids is None:
                item.issue_ids = []
            item.issue_ids.append(i)

    return content


def filter_meaningful_paragraphs(paragraphs: List[str], max_count: int = 100) -> List[dict]:
    """
    过滤出有意义的段落用于 LLM 审查

    Args:
        paragraphs: 段落列表，格式为 ["[0] 文本", "[1] 文本", ...]
        max_count: 最多返回的段落数

    Returns:
        过滤后的段落列表，格式为 [{"index": 0, "text": "文本"}, ...]
    """
    import re

    meaningful = []

    for p in paragraphs:
        # 解析索引和文件
        match = re.match(r'\[(\d+)\]\s*(.*)', p)
        if not match:
            continue

        index = int(match.group(1))
        text = match.group(2).strip()

        # 过滤规则
        # 1. 太短
        if len(text) < 10:
            continue

        # 2. 目录
        if re.search(r'\.{3,}|…{2,}', text):
            continue

        # 3. 纯数字和符号
        if re.match(r'^[\d\s\.\-—_，。、；：""''（）\(\)]+$', text):
            continue

        # 4. 页眉页脚特征
        if re.match(r'^第?\s*\d+\s*页', text) or re.match(r'^\d+\s*$', text):
            continue

        # 5. 标题行
        if len(text) < 20 and not re.search(r'[，。；：、]', text):
            continue

        meaningful.append({
            'index': index,
            'text': text
        })

        # 达到上限
        if len(meaningful) >= max_count:
            break

    return meaningful


def get_filtered_paragraphs_for_review(doc_content, max_count: int = 100) -> List[str]:
    """
    获取过滤后的段落列表（用于 LLM 审查）

    Args:
        doc_content: DocumentContent 对象
        max_count: 最多返回的段落数

    Returns:
        格式化的段落列表 ["[index] text", ...]
    """
    paragraphs = get_paragraphs_text(doc_content)
    filtered = filter_meaningful_paragraphs(paragraphs, max_count)

    # 转回原格式
    return [f"[{p['index']}] {p['text']}" for p in filtered]
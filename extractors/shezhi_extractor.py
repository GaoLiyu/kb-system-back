"""
涉执报告精确提取器
==================
针对涉执报告的表格结构精确提取
表格索引（基于分析结果）：
- 表格0/3/11: 结果汇总表
- 表格2: 权属表
- 表格5: 基础信息表
- 表格6: 因素描述表
- 表格7: 因素等级表
- 表格8: 因素指数表
- 表格9: 因素比率表
- 表格10: 修正计算表
"""

import os
import re
from typing import Dict, List, Optional
from docx import Document
from dataclasses import dataclass, field

from .table_utils import extract_property_rights_generic


@dataclass
class Position:
    """位置信息"""
    table_index: int = -1
    row_index: int = -1
    col_index: int = -1


@dataclass
class LocatedValue:
    """带位置的值"""
    value: any = None
    position: Position = field(default_factory=Position)
    raw_text: str = ""


@dataclass
class Factor:
    """因素数据"""
    name: str = ""
    description: str = ""
    level: str = ""
    index: int = 100
    desc_pos: Position = field(default_factory=Position)
    level_pos: Position = field(default_factory=Position)
    index_pos: Position = field(default_factory=Position)


@dataclass
class Case:
    """可比实例"""
    case_id: str = ""  # A/B/C
    address: LocatedValue = field(default_factory=LocatedValue)
    location: str = ""
    usage: str = ""
    data_source: str = ""
    transaction_price: LocatedValue = field(default_factory=LocatedValue)
    building_area: LocatedValue = field(default_factory=LocatedValue)
    transaction_date: str = ""
    
    # 新增字段
    district: str = ""           # 区域
    street: str = ""             # 街道/镇
    build_year: int = 0          # 建成年份
    total_floor: int = 0         # 总楼层
    current_floor: int = 0       # 所在楼层
    orientation: str = ""        # 朝向
    decoration: str = ""         # 装修
    structure: str = ""          # 结构
    
    # 修正系数
    transaction_correction: LocatedValue = field(default_factory=LocatedValue)
    market_correction: LocatedValue = field(default_factory=LocatedValue)
    location_correction: LocatedValue = field(default_factory=LocatedValue)
    physical_correction: LocatedValue = field(default_factory=LocatedValue)
    rights_correction: LocatedValue = field(default_factory=LocatedValue)
    adjusted_price: LocatedValue = field(default_factory=LocatedValue)
    
    # 因素
    location_factors: Dict[str, Factor] = field(default_factory=dict)
    physical_factors: Dict[str, Factor] = field(default_factory=dict)
    rights_factors: Dict[str, Factor] = field(default_factory=dict)


@dataclass
class Subject:
    """估价对象"""
    address: LocatedValue = field(default_factory=LocatedValue)
    building_area: LocatedValue = field(default_factory=LocatedValue)
    unit_price: LocatedValue = field(default_factory=LocatedValue)
    total_price: LocatedValue = field(default_factory=LocatedValue)
    
    # 权属
    # 权属
    cert_no: str = ""  # 房屋所有权证证号
    owner: str = ""  # 房屋所有权人
    structure: str = ""  # 结构
    floor: str = ""  # 楼层
    plan_usage: str = ""  # 规划用途
    land_use_type: str = ""  # 使用权类型
    land_type: str = ""  # 地类（用途）
    land_area: float = 0.0  # 土地面积
    end_date: str = ""  # 终止日期
    
    # 新增字段
    district: str = ""           # 区域（区/县）
    street: str = ""             # 街道/镇
    location: str = ""           # 位置
    data_source: str = ""        # 数据来源
    build_year: int = 0          # 建成年份
    total_floor: str = ""         # 总楼层
    current_floor: str = ""       # 所在楼层
    orientation: str = ""        # 朝向
    decoration: str = ""         # 装修状况
    land_end_date: str = ""      # 土地终止日期
    value_date: str = ""         # 价值时点
    usage: str = ""             # 用途
    appraisal_purpose: str = ""  # 估价目的

    # 新增价格内涵里面的字段
    property_scope: str = ""  # 财产范围
    payment_methods: str = ""  # 付款方式
    financing_conditions: str = ""  # 融资条件
    tax_burden: str = ""  # 税负
    unit_measurement: str = ""  # 计价单位
    price_type: str = ""  # 价格类型
    transaction_price: LocatedValue = field(default_factory=LocatedValue)
    transaction_date: str = ""
    
    # 因素
    location_factors: Dict[str, Factor] = field(default_factory=dict)
    physical_factors: Dict[str, Factor] = field(default_factory=dict)
    rights_factors: Dict[str, Factor] = field(default_factory=dict)


@dataclass
class ShezhiExtractionResult:
    """涉执报告提取结果"""
    source_file: str = ""
    subject: Subject = field(default_factory=Subject)
    cases: List[Case] = field(default_factory=list)
    
    # 最终结果
    final_unit_price: LocatedValue = field(default_factory=LocatedValue)
    final_total_price: LocatedValue = field(default_factory=LocatedValue)
    floor_factor: float = 1.0

    type: str = ""  # 类型


class ShezhiExtractor:
    """涉执报告提取器"""
    
    # 表格索引映射
    TABLE_RESULT_SUMMARY = 0      # 结果汇总
    TABLE_PROPERTY_RIGHTS = 2     # 权属表
    TABLE_BASIC_INFO = 5          # 基础信息
    TABLE_FACTOR_DESC = 6         # 因素描述
    TABLE_FACTOR_LEVEL = 7        # 因素等级
    TABLE_FACTOR_INDEX = 8        # 因素指数
    TABLE_FACTOR_RATIO = 9        # 因素比率
    TABLE_CORRECTION = 10         # 修正计算
    
    # 因素名称映射
    LOCATION_FACTORS = ['区域位置', '楼幢位置', '朝向', '交通条件', '配套设施', '环境质量', '景观', '物业管理']
    PHYSICAL_FACTORS = ['地形地势', '地质土壤', '开发程度', '建筑面积', '空间布局', '新旧程度', '装饰装修', '建筑结构', '物业类型', '设施设备']
    RIGHTS_FACTORS = ['规划条件', '土地使用期限', '担保物权设立', '租赁占用状况', '拖欠税费状况', '其他权益状况']
    
    def __init__(self, auto_detect: bool = False):
        self.doc = None
        self.tables = []
        self.full_text = ""
        self.auto_detect = auto_detect  # 是否自动检测表格索引

    def extract(self, doc_path: str) -> ShezhiExtractionResult:
        """提取涉执报告"""
        self.doc = Document(doc_path)
        self.tables = self.doc.tables
        self.full_text = "\n".join([p.text for p in self.doc.paragraphs])

        result = ShezhiExtractionResult(source_file=os.path.basename(doc_path), type='shezhi')

        print(f"\n📊 提取涉执报告: {os.path.basename(doc_path)}")
        print(f"   表格数量: {len(self.tables)}")

        # 自动检测表格索引（用于司法评估等变体）
        if self.auto_detect:
            self._auto_detect_table_indices()
            print(f"   ✓ 自动检测表格索引: 基础信息表={self.TABLE_BASIC_INFO}")

        # 1. 提取结果汇总
        self._extract_result_summary(result)
        print(f"   ✓ 结果汇总: {result.subject.address.value}")

        # 2. 提取权属信息
        self._extract_property_rights(result)
        print(f"   ✓ 权属信息: {result.subject.cert_no}")

        # 3. 提取基础信息
        self._extract_basic_info(result)
        print(f"   ✓ 基础信息: {len(result.cases)}个可比实例")

        # 4. 提取因素描述
        self._extract_factor_descriptions(result)

        # 5. 提取因素等级
        self._extract_factor_levels(result)

        # 6. 提取因素指数
        self._extract_factor_indices(result)
        print(f"   ✓ 因素数据: 描述/等级/指数")

        # 7. 提取修正系数
        self._extract_corrections(result)
        print(f"   ✓ 修正系数")

        # 8. 提取楼层修正系数
        self._extract_floor_factor(result)
        if result.floor_factor != 1.0:
            print(f"   ✓ 楼层修正: {result.floor_factor}")

        # 9. 提取扩展信息（建成年代、价值时点、估价目的等）
        self._extract_extended_info(result)

        # 10. 解析区域信息
        self._parse_district(result)

        return result

    def _auto_detect_table_indices(self):
        """
        自动检测关键表格索引（涉执报告）
        采用打分制：对每个表抽取前几行/列做文本块，然后按关键词规则打分，分别选最高分的表。
        """

        def norm(s: str) -> str:
            if not s:
                return ""
            return (
                s.replace("\u3000", " ")
                .replace("\n", " ")
                .replace("\t", " ")
                .strip()
            )

        def table_block(table, max_rows=8, max_cols=12) -> str:
            parts = []
            rN = min(len(table.rows), max_rows)
            for r in range(rN):
                row = table.rows[r]
                cN = min(len(row.cells), max_cols)
                for c in range(cN):
                    parts.append(norm(row.cells[c].text))
            return " ".join([p for p in parts if p])

        def compact(s: str) -> str:
            return s.replace(" ", "")

        def count_hits(text: str, keys) -> int:
            return sum(1 for k in keys if k in text)

        def has_any(text: str, keys) -> bool:
            return any(k in text for k in keys)

        # best_scores: name -> (score, index)
        best = {
            "result": (-1, self.TABLE_RESULT_SUMMARY),
            "rights": (-1, self.TABLE_PROPERTY_RIGHTS),
            "basic": (-1, self.TABLE_BASIC_INFO),
            "desc": (-1, self.TABLE_FACTOR_DESC),
            "level": (-1, self.TABLE_FACTOR_LEVEL),
            "index": (-1, self.TABLE_FACTOR_INDEX),
            "ratio": (-1, self.TABLE_FACTOR_RATIO),
            "corr": (-1, self.TABLE_CORRECTION),
        }

        # 因素表通用的 A/B/C 头
        def has_abc_header(t: str) -> bool:
            t2 = compact(t)
            return ("估价对象" in t2) and ("可比实例A" in t2) and ("可比实例B" in t2) and ("可比实例C" in t2)

        for i, table in enumerate(self.tables):
            if len(table.rows) == 0:
                continue

            rows = len(table.rows)
            cols = len(table.columns) if table.columns else 0

            block = table_block(table, max_rows=10, max_cols=12)
            t = compact(block)

            # ---------------- 1) 结果汇总表 ----------------
            # 特征：同时出现“单价”“总价”，且常见“建筑面积”“元/平方米”等
            score_result = 0
            if ("单价" in t and "总价" in t) or ("元/平方米" in t and "总价" in t):
                score_result += 10
            score_result += count_hits(t, ["建筑面积", "平方米", "估价对象坐落", "万元", "大写"])
            # 通常汇总表较短，但不写死，只加分
            if rows <= 6:
                score_result += 2
            if score_result > best["result"][0]:
                best["result"] = (score_result, i)

            # ---------------- 2) 权属表 ----------------
            # 特征：不动产权证/不动产权利人/结构/规划用途/土地面积/终止日期 等
            score_rights = 0
            rights_strong = ["不动产权证", "不动产权利人", "不动产权属", "权属登记"]
            rights_weak = ["坐落", "结构", "规划用途", "使用权类型", "地类", "土地面积", "终止日期", "建筑面积"]
            score_rights += 3 * count_hits(t, rights_strong)
            score_rights += 1 * count_hits(t, rights_weak)
            if score_rights > best["rights"][0]:
                best["rights"] = (score_rights, i)

            # ---------------- 3) 基础信息表 ----------------
            # 特征：项目/估价对象/可比实例(至少A/B/C) + 行里有“地址/用途/来源/交易日期/建筑面积”等
            score_basic = 0
            if ("项目" in t and "估价对象" in t and "可比实例" in t):
                score_basic += 8
            # 允许模板是“项目 + 估价对象 + 可比实例A/B/C”
            if ("可比实例A" in t and "可比实例B" in t and "可比实例C" in t):
                score_basic += 4
            score_basic += count_hits(t, ["地址", "坐落", "用途", "来源", "交易日期", "建筑面积", "成交", "价格类型",
                                          "财产范围"])
            # 基础信息表通常行数较多
            if rows >= 10:
                score_basic += 2
            if score_basic > best["basic"][0]:
                best["basic"] = (score_basic, i)

            # ---------------- 4) 因素类表（描述/等级/指数/比率）----------------
            # 四张表共同特点：ABC 列头 + 内容出现 “区位/实物/权益” 等分类项
            if has_abc_header(block) or has_any(t, ["区位状况", "实物状况", "权益状况"]):
                # 4.1 描述表：出现“描述性”的词更密集（如“状况/条件/质量/配套/装修”等）
                score_desc = 0
                if has_any(t, ["区位状况", "实物状况", "权益状况"]):
                    score_desc += 2
                score_desc += count_hits(t, ["交通条件", "配套设施", "环境质量", "物业管理", "装饰装修", "建筑结构",
                                             "新旧程度"])
                # 描述表一般不会大量出现“指数/比率”
                score_desc -= 2 * count_hits(t, ["指数", "比率"])
                if score_desc > best["desc"][0]:
                    best["desc"] = (score_desc, i)

                # 4.2 等级表：关键词“等级/优/良/中/差/较优/一般”等
                score_level = 0
                score_level += 3 * count_hits(t, ["等级"])
                score_level += count_hits(t, ["优", "良", "中", "差", "较优", "一般"])
                # 等级表一般不出现“指数/比率/100”密集
                score_level -= count_hits(t, ["指数", "比率", "100"])
                if score_level > best["level"][0]:
                    best["level"] = (score_level, i)

                # 4.3 指数表：关键词“指数/100/95/105”等数字密集
                score_index = 0
                score_index += 4 * count_hits(t, ["指数"])
                # 常见指数基准 100（不写死，只加分）
                if "100" in t:
                    score_index += 2
                # 指数表一般不出现“比率/系数(%)”密集
                score_index -= 2 * count_hits(t, ["比率"])
                if score_index > best["index"][0]:
                    best["index"] = (score_index, i)

                # 4.4 比率表：关键词“比率/系数/%/修正系数/比准”等
                score_ratio = 0
                score_ratio += 4 * count_hits(t, ["比率"])
                score_ratio += 2 * count_hits(t, ["系数", "%", "修正"])
                if score_ratio > best["ratio"][0]:
                    best["ratio"] = (score_ratio, i)

            # ---------------- 5) 修正计算表 ----------------
            # 特征：修正系数/修正结果/比准价格/P1.. 等（涉执模板里常见交易情况/日期/区位/实物/权益修正）
            score_corr = 0
            corr_strong = ["修正", "比准", "修正结果", "比准价格"]
            corr_weak = ["交易情况", "交易日期", "区位", "实物", "权益", "调整", "系数"]
            score_corr += 2 * count_hits(t, corr_strong)
            score_corr += 1 * count_hits(t, corr_weak)
            # 修正表一般更大些
            if rows >= 8 and cols >= 5:
                score_corr += 2
            if score_corr > best["corr"][0]:
                best["corr"] = (score_corr, i)

        # --------- 落盘：设置表索引（设一个阈值，避免误判）---------
        # 阈值可以按你后续样本再调，这里先给一个保守值
        if best["result"][0] >= 8:
            self.TABLE_RESULT_SUMMARY = best["result"][1]
        if best["rights"][0] >= 6:
            self.TABLE_PROPERTY_RIGHTS = best["rights"][1]
        if best["basic"][0] >= 8:
            self.TABLE_BASIC_INFO = best["basic"][1]
        if best["corr"][0] >= 6:
            self.TABLE_CORRECTION = best["corr"][1]

        # 因素四张表：如果都识别到了，用识别结果；否则退化为“从基础信息表之后按顺序猜”
        found_desc = best["desc"][0] >= 3
        found_level = best["level"][0] >= 3
        found_index = best["index"][0] >= 3
        found_ratio = best["ratio"][0] >= 3

        if found_desc:
            self.TABLE_FACTOR_DESC = best["desc"][1]
        if found_level:
            self.TABLE_FACTOR_LEVEL = best["level"][1]
        if found_index:
            self.TABLE_FACTOR_INDEX = best["index"][1]
        if found_ratio:
            self.TABLE_FACTOR_RATIO = best["ratio"][1]

        # 兜底策略：当因素表没全识别出来时，用“基础信息表 + 偏移”兜底，但不强依赖
        if not (found_desc and found_level and found_index and found_ratio):
            base = self.TABLE_BASIC_INFO
            # 只有当 base 合理时才兜底
            if 0 <= base < len(self.tables):
                # 兜底偏移（与你原来的逻辑一致，但仅作为 fallback）
                if not found_desc:
                    self.TABLE_FACTOR_DESC = min(base + 1, len(self.tables) - 1)
                if not found_level:
                    self.TABLE_FACTOR_LEVEL = min(base + 2, len(self.tables) - 1)
                if not found_index:
                    self.TABLE_FACTOR_INDEX = min(base + 3, len(self.tables) - 1)
                if not found_ratio:
                    self.TABLE_FACTOR_RATIO = min(base + 4, len(self.tables) - 1)
                # 修正表也兜底一下（前面已识别则不覆盖）
                if best["corr"][0] < 6:
                    self.TABLE_CORRECTION = min(base + 5, len(self.tables) - 1)

    def _get_cell_value(self, table_idx: int, row_idx: int, col_idx: int) -> LocatedValue:
        """获取单元格值（带位置）"""
        try:
            table = self.tables[table_idx]
            cell = table.rows[row_idx].cells[col_idx]
            return LocatedValue(
                value=cell.text.strip(),
                position=Position(table_idx, row_idx, col_idx),
                raw_text=cell.text.strip()
            )
        except:
            return LocatedValue()
    
    def _extract_result_summary(self, result: ShezhiExtractionResult):
        """提取结果汇总表"""
        table = self.tables[self.TABLE_RESULT_SUMMARY]
        
        # 第二行是数据行
        if len(table.rows) >= 2:
            row = table.rows[1]
            cells = [c.text.strip() for c in row.cells]
            
            result.subject.address = LocatedValue(
                value=cells[0] if cells else "",
                position=Position(self.TABLE_RESULT_SUMMARY, 1, 0),
                raw_text=cells[0] if cells else ""
            )
            
            if len(cells) >= 2:
                try:
                    result.subject.building_area = LocatedValue(
                        value=float(cells[1]),
                        position=Position(self.TABLE_RESULT_SUMMARY, 1, 1),
                        raw_text=cells[1]
                    )
                except:
                    pass
            
            if len(cells) >= 3:
                try:
                    result.subject.unit_price = LocatedValue(
                        value=float(cells[2]),
                        position=Position(self.TABLE_RESULT_SUMMARY, 1, 2),
                        raw_text=cells[2]
                    )
                    result.final_unit_price = result.subject.unit_price
                except:
                    pass
            
            if len(cells) >= 4:
                # 提取总价数字
                total_text = cells[3]
                match = re.search(r'([\d.]+)', total_text)
                if match:
                    result.subject.total_price = LocatedValue(
                        value=float(match.group(1)),
                        position=Position(self.TABLE_RESULT_SUMMARY, 1, 3),
                        raw_text=total_text
                    )
                    result.final_total_price = result.subject.total_price

    def _extract_property_rights(self, result: ShezhiExtractionResult):
        """提取权属表（使用 table_utils：表头定位 + 列映射，不改变你的结果结构）"""
        if len(self.tables) <= self.TABLE_PROPERTY_RIGHTS:
            return

        table = self.tables[self.TABLE_PROPERTY_RIGHTS]

        # 这个 setter 只负责把 utils 识别到的字段写回你现有 Subject 字段
        def subject_setter(key: str, value):
            # 注意：这里完全不改你的字段名/类型，只是赋值方式不同
            if key == "cert_no":
                if value and not result.subject.cert_no:
                    result.subject.cert_no = str(value).strip()

            elif key == "owner":
                if value and not result.subject.owner:
                    result.subject.owner = str(value).strip()

            elif key == "address":
                # shezhi 的 Subject.address 是 LocatedValue
                if value and not result.subject.address.value:
                    result.subject.address.value = str(value).strip()

            elif key == "structure":
                if value and not result.subject.structure:
                    result.subject.structure = str(value).strip()

            elif key == "floor":
                if value and not result.subject.floor:
                    result.subject.floor = str(value).strip()

            elif key == "plan_usage":
                if value and not result.subject.plan_usage:
                    result.subject.plan_usage = str(value).strip()

            # ---- 土地块（shezhi Subject 里是 land_use_type / land_type / land_area / end_date）----
            elif key == "land_use_type":
                if value and not result.subject.land_use_type:
                    result.subject.land_use_type = str(value).strip()

            elif key == "land_type":
                if value and not result.subject.land_type:
                    result.subject.land_type = str(value).strip()

            elif key == "land_area":
                # utils 解析出来一般是 float
                if value is not None and (not result.subject.land_area):
                    try:
                        result.subject.land_area = float(value)
                    except:
                        pass

            elif key == "end_date":
                if value and not result.subject.end_date:
                    result.subject.end_date = str(value).strip()

            # 有些权属表里会有土地证号/土地权利人等字段，但你 shezhi 的 Subject 没定义这些字段
            # 所以这里故意忽略，不会改变你的数据结构

        # detect_land=True：让 utils 试着解析土地块（如果表里有）
        extract_property_rights_generic(table, subject_setter=subject_setter, detect_land=True)

    def _extract_basic_info(self, result: ShezhiExtractionResult):
        """提取基础信息表"""
        table = self.tables[self.TABLE_BASIC_INFO]
        
        # 初始化三个可比实例
        result.cases = [Case(case_id='A'), Case(case_id='B'), Case(case_id='C')]
        
        # 列映射: 估价对象=2, A=3, B=4, C=5
        COL_SUBJECT = 2
        COL_A = 3
        COL_B = 4
        COL_C = 5
        
        for row_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            
            if len(cells) < 6:
                continue
            
            # 获取行标签（前两列可能合并）
            label = cells[0] + cells[1] if len(cells) > 1 else cells[0]
            label = label.replace(' ', '').replace('\u3000', '')
            
            if '地址' in label or '坐落' in label:
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.address = LocatedValue(
                            value=cells[col],
                            position=Position(self.TABLE_BASIC_INFO, row_idx, col),
                            raw_text=cells[col]
                        )
            
            elif '位置' in label and '楼' not in label:
                if COL_SUBJECT < len(cells):
                    result.subject.location = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.location = cells[col]
            
            elif '来源' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.data_source = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.data_source = cells[col]
            
            elif '用途' in label:
                # 估价对象用途
                if COL_SUBJECT < len(cells):
                    result.subject.usage = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.usage = cells[col]

            elif '财产范围' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.property_scope = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.property_scope = cells[col]

            elif '付款方式' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.payment_methods = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.payment_methods = cells[col]

            elif '融资条件' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.financing_conditions = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.financing_conditions = cells[col]

            elif '税' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.tax_burden = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.tax_burden = cells[col]

            elif '单位' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.unit_measurement = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.unit_measurement = cells[col]

            elif '价格类型' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.price_type = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.price_type = cells[col]

            elif '评估面积' in label or '建筑面积' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.building_area = LocatedValue(
                        value=cells[COL_SUBJECT],
                        position=Position(self.TABLE_BASIC_INFO, row_idx, COL_SUBJECT),
                        raw_text=cells[COL_SUBJECT]
                    )
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        try:
                            area = float(re.sub(r'[^\d.]', '', cells[col]))
                            case.building_area = LocatedValue(
                                value=area,
                                position=Position(self.TABLE_BASIC_INFO, row_idx, col),
                                raw_text=cells[col]
                            )
                        except:
                            pass

            elif '交易日期' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.transaction_date = cells[COL_SUBJECT]
                # result.subject.transaction_date = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.transaction_date = cells[col]
            
            elif '成交基价' in label or '交易价格' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.transaction_price = LocatedValue(
                        value=cells[COL_SUBJECT],
                        position=Position(self.TABLE_BASIC_INFO, row_idx, COL_SUBJECT),
                        raw_text=cells[COL_SUBJECT]
                    )
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        try:
                            price = float(re.sub(r'[^\d.]', '', cells[col]))
                            case.transaction_price = LocatedValue(
                                value=price,
                                position=Position(self.TABLE_BASIC_INFO, row_idx, col),
                                raw_text=cells[col]
                            )
                        except:
                            pass
    
    def _extract_factor_descriptions(self, result: ShezhiExtractionResult):
        """提取因素描述表（按固定列读取，避免去重导致列错位）

        表6固定为6列（基于该模板分析）：
        0=分类, 1=因素名称, 2=估价对象, 3/4/5=可比实例A/B/C
        """
        table = self.tables[self.TABLE_FACTOR_DESC]

        COL_CATEGORY = 0
        COL_FACTOR = 1
        COL_SUBJECT = 2
        COL_A = 3
        COL_B = 4
        COL_C = 5

        category_alias = {
            '区位状况': '区位状况',
            '实物状况': '实物状况',
            '实物因素': '实物状况',
            '权益状况': '权益状况',
            '权益因素': '权益状况',
        }

        current_category = ""
        for row_idx, row in enumerate(table.rows[1:], 1):  # 跳过表头
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            if len(cells) < 6:
                continue

            raw_category = (cells[COL_CATEGORY] or '').replace(' ', '').replace('\u3000', '').replace('　', '')
            factor_name = (cells[COL_FACTOR] or '').replace(' ', '').replace('\u3000', '').replace('　', '')

            # 跳过交易类
            if raw_category in ('交易情况', '交易日期') or factor_name in ('交易情况', '交易日期'):
                continue

            # 分类更新
            if raw_category in category_alias:
                current_category = category_alias[raw_category]

            if not factor_name:
                continue

            # 确定因素类别
            if factor_name in self.LOCATION_FACTORS or current_category == '区位状况':
                factor_type = 'location'
            elif factor_name in self.PHYSICAL_FACTORS or current_category == '实物状况':
                factor_type = 'physical'
            elif factor_name in self.RIGHTS_FACTORS or current_category == '权益状况':
                factor_type = 'rights'
            else:
                continue

            factor_key = self._normalize_factor_name(factor_name)

            # 估价对象
            subject_value = cells[COL_SUBJECT]
            if subject_value != '':
                if factor_type == 'location':
                    d = result.subject.location_factors
                elif factor_type == 'physical':
                    d = result.subject.physical_factors
                else:
                    d = result.subject.rights_factors
                f = d.get(factor_key) or Factor(name=factor_key)
                f.description = subject_value
                f.desc_pos = Position(self.TABLE_FACTOR_DESC, row_idx, COL_SUBJECT)
                d[factor_key] = f
                self._sync_subject_fields_from_factor(result.subject, factor_key, subject_value)

            # 可比实例A/B/C
            for i, case in enumerate(result.cases):
                col = [COL_A, COL_B, COL_C][i]
                value = cells[col]
                if value == '':
                    continue

                if factor_type == 'location':
                    d = case.location_factors
                elif factor_type == 'physical':
                    d = case.physical_factors
                else:
                    d = case.rights_factors

                f = d.get(factor_key) or Factor(name=factor_key)
                f.description = value
                f.desc_pos = Position(self.TABLE_FACTOR_DESC, row_idx, col)
                d[factor_key] = f
                self._sync_case_fields_from_factor(case, factor_key, value)

    def _extract_factor_levels(self, result: ShezhiExtractionResult):
        """提取因素等级表（按固定列读取，避免去重导致列错位）"""
        table = self.tables[self.TABLE_FACTOR_LEVEL]

        COL_CATEGORY = 0
        COL_FACTOR = 1
        COL_SUBJECT = 2
        COL_A = 3
        COL_B = 4
        COL_C = 5

        category_alias = {
            '区位状况': '区位状况',
            '实物状况': '实物状况',
            '实物因素': '实物状况',
            '权益状况': '权益状况',
            '权益因素': '权益状况',
        }

        current_category = ""
        for row_idx, row in enumerate(table.rows[1:], 1):
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            if len(cells) < 6:
                continue

            raw_category = (cells[COL_CATEGORY] or '').replace(' ', '').replace('\u3000', '').replace('　', '')
            factor_name = (cells[COL_FACTOR] or '').replace(' ', '').replace('\u3000', '').replace('　', '')

            if raw_category in ('交易情况', '交易日期') or factor_name in ('交易情况', '交易日期'):
                continue

            if raw_category in category_alias:
                current_category = category_alias[raw_category]

            if not factor_name:
                continue

            if factor_name in self.LOCATION_FACTORS or current_category == '区位状况':
                factor_type = 'location'
            elif factor_name in self.PHYSICAL_FACTORS or current_category == '实物状况':
                factor_type = 'physical'
            elif factor_name in self.RIGHTS_FACTORS or current_category == '权益状况':
                factor_type = 'rights'
            else:
                continue

            factor_key = self._normalize_factor_name(factor_name)

            # 估价对象
            subject_value = cells[COL_SUBJECT]
            if subject_value != '':
                if factor_type == 'location':
                    d = result.subject.location_factors
                elif factor_type == 'physical':
                    d = result.subject.physical_factors
                else:
                    d = result.subject.rights_factors
                f = d.get(factor_key) or Factor(name=factor_key)
                f.level = subject_value
                f.level_pos = Position(self.TABLE_FACTOR_LEVEL, row_idx, COL_SUBJECT)
                d[factor_key] = f

            # 可比实例A/B/C
            for i, case in enumerate(result.cases):
                col = [COL_A, COL_B, COL_C][i]
                value = cells[col]
                if value == '':
                    continue
                if factor_type == 'location':
                    d = case.location_factors
                elif factor_type == 'physical':
                    d = case.physical_factors
                else:
                    d = case.rights_factors
                f = d.get(factor_key) or Factor(name=factor_key)
                f.level = value
                f.level_pos = Position(self.TABLE_FACTOR_LEVEL, row_idx, col)
                d[factor_key] = f

    def _extract_factor_indices(self, result: ShezhiExtractionResult):
        """提取因素指数表（按固定列读取，避免去重导致列错位）"""
        table = self.tables[self.TABLE_FACTOR_INDEX]

        COL_CATEGORY = 0
        COL_FACTOR = 1
        COL_SUBJECT = 2
        COL_A = 3
        COL_B = 4
        COL_C = 5

        category_alias = {
            '区位状况': '区位状况',
            '实物状况': '实物状况',
            '实物因素': '实物状况',
            '权益状况': '权益状况',
            '权益因素': '权益状况',
        }

        def to_int(v: str) -> int:
            try:
                return int(re.sub(r'[^0-9]', '', v))
            except Exception:
                return 100

        current_category = ""
        for row_idx, row in enumerate(table.rows[1:], 1):
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            if len(cells) < 6:
                continue

            raw_category = (cells[COL_CATEGORY] or '').replace(' ', '').replace('\u3000', '').replace('　', '')
            factor_name = (cells[COL_FACTOR] or '').replace(' ', '').replace('\u3000', '').replace('　', '')

            if raw_category in ('交易情况', '交易日期') or factor_name in ('交易情况', '交易日期'):
                continue

            if raw_category in category_alias:
                current_category = category_alias[raw_category]

            if not factor_name:
                continue

            if factor_name in self.LOCATION_FACTORS or current_category == '区位状况':
                factor_type = 'location'
            elif factor_name in self.PHYSICAL_FACTORS or current_category == '实物状况':
                factor_type = 'physical'
            elif factor_name in self.RIGHTS_FACTORS or current_category == '权益状况':
                factor_type = 'rights'
            else:
                continue

            factor_key = self._normalize_factor_name(factor_name)

            # 估价对象
            subject_value = cells[COL_SUBJECT]
            if subject_value != '':
                if factor_type == 'location':
                    d = result.subject.location_factors
                elif factor_type == 'physical':
                    d = result.subject.physical_factors
                else:
                    d = result.subject.rights_factors
                f = d.get(factor_key) or Factor(name=factor_key)
                f.index = to_int(subject_value)
                f.index_pos = Position(self.TABLE_FACTOR_INDEX, row_idx, COL_SUBJECT)
                d[factor_key] = f

            # 可比实例A/B/C
            for i, case in enumerate(result.cases):
                col = [COL_A, COL_B, COL_C][i]
                value = cells[col]
                if value == '':
                    continue
                if factor_type == 'location':
                    d = case.location_factors
                elif factor_type == 'physical':
                    d = case.physical_factors
                else:
                    d = case.rights_factors
                f = d.get(factor_key) or Factor(name=factor_key)
                f.index = to_int(value)
                f.index_pos = Position(self.TABLE_FACTOR_INDEX, row_idx, col)
                d[factor_key] = f

    def _extract_corrections(self, result: ShezhiExtractionResult):
        """提取修正系数"""
        table = self.tables[self.TABLE_CORRECTION]

        # 修正计算表列: A=1, B=2, C=3
        COL_A = 1

        ROW_MAPPING = {
            '交易价格': 'transaction_price',
            '交易情况修正': 'transaction_correction',
            '市场状况': 'market_correction',
            '区位状况': 'location_correction',
            '实物状况': 'physical_correction',
            '权益状况': 'rights_correction',
            '修正后单价': 'adjusted_price',
        }

        for row_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]

            if len(cells) < 2:
                continue

            label = cells[0].replace(' ', '').replace('\u3000', '')

            field_name = None
            for key, field in ROW_MAPPING.items():
                if key in label:
                    field_name = field
                    break

            if not field_name:
                continue

            for i, case in enumerate(result.cases):
                col = COL_A + i
                if col < len(cells):
                    try:
                        value = float(cells[col])
                        loc_val = LocatedValue(
                            value=value,
                            position=Position(self.TABLE_CORRECTION, row_idx, col),
                            raw_text=cells[col]
                        )
                        setattr(case, field_name, loc_val)
                    except:
                        pass

    def _extract_floor_factor(self, result: ShezhiExtractionResult):
        """提取楼层修正系数"""
        match = re.search(r'×\s*(\d+)%\s*[＝=]', self.full_text)
        if match:
            result.floor_factor = int(match.group(1)) / 100

    def _extract_extended_info(self, result: ShezhiExtractionResult):
        """提取扩展信息（建成年代、价值时点、估价目的等）"""

        # 1. 建成年代 - 从段落文本中提取
        # 匹配模式: "建成于XXXX年" 或 "约建成于本世纪初" 或 "建成年代：XXXX"
        build_patterns = [
            r'建成于(\d{4})年',
            r'约(\d{4})年建成',
            r'建成年代[：:]\s*(\d{4})',
            r'(\d{4})年建成',
            r'约建成于本世纪初',  # 特殊处理
            r'建成于上世纪(\d{2})年代',
        ]

        for pattern in build_patterns:
            match = re.search(pattern, self.full_text)
            if match:
                if '本世纪初' in pattern:
                    result.subject.build_year = 2000
                elif len(match.groups()) > 0:
                    year_str = match.group(1)
                    if len(year_str) == 2:
                        # 处理"90年代"这种格式
                        result.subject.build_year = 1900 + int(year_str)
                    else:
                        result.subject.build_year = int(year_str)
                break

        # 2. 价值时点 - 从段落文本中提取
        value_date_patterns = [
            r'价值时点[：:]\s*(\d{4})[年\.](\d{1,2})[月\.](\d{1,2})',
            r'价值时点(\d{4})\.(\d{1,2})\.(\d{1,2})',
            r'价值时点为(\d{4})年(\d{1,2})月(\d{1,2})日',
        ]

        for pattern in value_date_patterns:
            match = re.search(pattern, self.full_text)
            if match:
                result.subject.value_date = f"{match.group(1)}-{match.group(2).zfill(2)}-{match.group(3).zfill(2)}"
                break

        # 3. 估价目的 - 从段落文本中提取
        purpose_patterns = [
            r'估价目的[：:是为]*(.{5,50}?)(?:。|$)',
            r'本次估价目的是(.{5,50}?)(?:。|$)',
        ]

        for pattern in purpose_patterns:
            match = re.search(pattern, self.full_text)
            if match:
                result.subject.appraisal_purpose = match.group(1).strip()
                break

        # 4. 土地终止日期 - 从权属表中提取
        if len(self.tables) > self.TABLE_PROPERTY_RIGHTS:
            table = self.tables[self.TABLE_PROPERTY_RIGHTS]
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                for i, cell in enumerate(cells):
                    if '终止' in cell and i + 1 < len(cells):
                        # 找下一行同一列
                        pass
                    # 匹配日期格式 YYYY/MM/DD
                    date_match = re.search(r'(\d{4}/\d{1,2}/\d{1,2})', cell)
                    if date_match:
                        result.subject.land_end_date = date_match.group(1)

        # 5. 解析楼层信息（从字符串"8/10"解析为数字）
        if result.subject.floor and '/' in result.subject.floor:
            parts = result.subject.floor.split('/')
            if len(parts) == 2:
                try:
                    result.subject.current_floor = str(parts[0])
                    result.subject.total_floor = str(parts[1])
                except:
                    pass

    def _parse_district(self, result: ShezhiExtractionResult):
        """从地址解析区域信息"""
        address = result.subject.address.value or ""

        # 常见区域关键词
        district_patterns = [
            r'([\u4e00-\u9fa5]{2,4}区)',   # XX区
            r'([\u4e00-\u9fa5]{2,4}县)',   # XX县
            r'([\u4e00-\u9fa5]{2,4}市)',   # XX市（县级市）
        ]

        for pattern in district_patterns:
            match = re.search(pattern, address)
            if match:
                result.subject.district = match.group(1)
                break

        # 街道/镇
        street_patterns = [
            r'([\u4e00-\u9fa5]{2,6}街道)',
            r'([\u4e00-\u9fa5]{2,4}镇)',
            r'([\u4e00-\u9fa5]{2,4}乡)',
        ]

        for pattern in street_patterns:
            match = re.search(pattern, address)
            if match:
                result.subject.street = match.group(1)
                break

        # 同样处理可比实例
        for case in result.cases:
            case_addr = case.address.value or ""

            for pattern in district_patterns:
                match = re.search(pattern, case_addr)
                if match:
                    case.district = match.group(1)
                    break

            for pattern in street_patterns:
                match = re.search(pattern, case_addr)
                if match:
                    case.street = match.group(1)
                    break

    def _sync_subject_fields_from_factor(self, subject: Subject, factor_key: str, val: str):
        """把因素描述同步到 Subject 的新增字段（不改变输出结构）"""
        v = (val or "").strip()
        if not v:
            return
        if factor_key == "orientation" and not subject.orientation:
            subject.orientation = v
        elif factor_key == "decoration" and not subject.decoration:
            subject.decoration = v
        elif factor_key == "structure" and not subject.structure:
            subject.structure = v

    def _sync_case_fields_from_factor(self, case: Case, factor_key: str, val: str):
        """把因素描述同步到 Case 的新增字段（不改变输出结构）"""
        v = (val or "").strip()
        if not v:
            return
        if factor_key == "orientation" and not case.orientation:
            case.orientation = v
        elif factor_key == "decoration" and not case.decoration:
            case.decoration = v
        elif factor_key == "structure" and not case.structure:
            case.structure = v


    def _normalize_factor_name(self, name: str) -> str:
        """标准化因素名称"""
        name = name.replace(' ', '').replace('\u3000', '').replace('　', '')

        mapping = {
            '区域位置': 'location_region',
            '楼幢位置': 'location_building',
            '朝向': 'orientation',
            '交通条件': 'traffic',
            '配套设施': 'facilities',
            '环境质量': 'environment',
            '景观': 'landscape',
            '物业管理': 'property_management',
            '地形地势': 'terrain',
            '地质土壤': 'geology',
            '开发程度': 'development',
            '建筑面积': 'area',
            '空间布局': 'layout',
            '新旧程度': 'age',
            '装饰装修': 'decoration',
            '建筑结构': 'structure',
            '物业类型': 'property_type',
            '设施设备': 'equipment',
            '规划条件': 'planning',
            '土地使用期限': 'land_term',
            '担保物权设立': 'mortgage',
            '租赁占用状况': 'lease',
            '拖欠税费状况': 'tax',
            '其他权益状况': 'other_rights',
        }

        return mapping.get(name, name)


# ============================================================================
# 测试
# ============================================================================

if __name__ == "__main__":
    extractor = ShezhiExtractor()
    result = extractor.extract("./data/docs/涉执报告-比较法.docx")

    print(result.subject)

"""
批量评估报告提取器（市场价值-现状价值）
====================================
针对批量评估报告的表格结构精确提取
特点：一个文档包含多个估价对象（如23套房产）

表格结构：
- 表格0: 批量汇总表（序号、坐落、建筑面积、评估总价）
- 表格4-9: 第一组可比实例
- 表格10: 楼层修正结果表
- 表格11+: 更多可比实例组...
"""

import os
import re
from typing import Dict, List, Optional
from docx import Document
from dataclasses import dataclass, field


@dataclass
class Position:
    """位置信息"""
    table_index: int = -1
    row_index: int = -1
    col_index: int = -1


@dataclass
class LocatedValue:
    """带位置的值"""
    value: any = None
    position: Position = field(default_factory=Position)
    raw_text: str = ""


@dataclass
class Factor:
    """因素数据"""
    name: str = ""
    description: str = ""
    level: str = ""
    index: int = 100


@dataclass
class Case:
    """可比实例"""
    case_id: str = ""  # A/B/C
    address: LocatedValue = field(default_factory=LocatedValue)
    location: str = ""
    usage: str = ""
    data_source: str = ""
    transaction_price: LocatedValue = field(default_factory=LocatedValue)
    building_area: LocatedValue = field(default_factory=LocatedValue)
    transaction_date: str = ""

    district: str = ""
    street: str = ""
    build_year: int = 0
    total_floor: int = 0
    current_floor: int = 0
    orientation: str = ""
    decoration: str = ""
    structure: str = ""

    # 修正系数
    transaction_correction: LocatedValue = field(default_factory=LocatedValue)
    market_correction: LocatedValue = field(default_factory=LocatedValue)
    location_correction: LocatedValue = field(default_factory=LocatedValue)
    physical_correction: LocatedValue = field(default_factory=LocatedValue)
    rights_correction: LocatedValue = field(default_factory=LocatedValue)
    adjusted_price: LocatedValue = field(default_factory=LocatedValue)

    # 因素
    location_factors: Dict[str, Factor] = field(default_factory=dict)
    physical_factors: Dict[str, Factor] = field(default_factory=dict)
    rights_factors: Dict[str, Factor] = field(default_factory=dict)


@dataclass
class BatchSubject:
    """批量估价对象"""
    seq_no: int = 0  # 序号
    address: str = ""  # 坐落
    building_area: float = 0  # 建筑面积
    total_price: float = 0  # 评估总价（万元）
    unit_price: float = 0  # 评估单价（元/㎡）
    floor_factor: float = 1.0  # 楼层系数

    # 可选的详细字段
    current_floor: int = 0
    total_floor: int = 0


@dataclass
class XianzhibExtractionResult:
    """批量评估报告提取结果"""
    source_file: str = ""

    # 批量估价对象列表
    subjects: List[BatchSubject] = field(default_factory=list)

    # 可比实例（可能有多组）
    case_groups: List[List[Case]] = field(default_factory=list)

    # 汇总信息
    total_count: int = 0
    total_area: float = 0
    total_value: float = 0

    # 基准价（用于计算各套房产的价格）
    base_price: float = 0


class XianzhibExtractor:
    """批量评估报告提取器"""

    # 因素名称映射
    LOCATION_FACTORS = ['区域位置', '楼幢位置', '朝向', '交通条件', '配套设施', '环境质量', '景观', '物业管理']
    PHYSICAL_FACTORS = ['地形地势', '地质土壤', '开发程度', '建筑面积', '空间布局', '新旧程度', '装饰装修', '建筑结构',
                        '物业类型', '设施设备']
    RIGHTS_FACTORS = ['规划条件', '土地使用期限', '担保物权设立', '租赁占用状况', '拖欠税费状况', '其他权益状况']

    def __init__(self):
        self.doc = None
        self.tables = []
        self.full_text = ""

    def extract(self, doc_path: str) -> XianzhibExtractionResult:
        """提取批量评估报告"""
        self.doc = Document(doc_path)
        self.tables = self.doc.tables
        self.full_text = "\n".join([p.text for p in self.doc.paragraphs])

        result = XianzhibExtractionResult(source_file=os.path.basename(doc_path))

        print(f"\n📊 提取批量评估报告: {os.path.basename(doc_path)}")
        print(f"   表格数量: {len(self.tables)}")

        # 1. 提取批量汇总表
        self._extract_batch_summary(result)
        print(f"   ✓ 批量汇总: {result.total_count}个估价对象")

        # 2. 查找并提取可比实例组
        self._extract_case_groups(result)
        print(f"   ✓ 可比实例组: {len(result.case_groups)}组")

        # 3. 提取楼层修正信息
        self._extract_floor_corrections(result)

        # 4. 计算汇总信息
        result.total_area = sum(s.building_area for s in result.subjects)
        result.total_value = sum(s.total_price for s in result.subjects)

        return result

    def _extract_batch_summary(self, result: XianzhibExtractionResult):
        """提取批量汇总表"""
        # 表格0应该是批量汇总表
        if len(self.tables) == 0:
            return

        table = self.tables[0]

        # 检查表头确认是批量汇总表
        if len(table.rows) > 0:
            header = [c.text.strip() for c in table.rows[0].cells]
            # 应该包含：序号、坐落、建筑面积、评估总价
            if not any('序号' in h for h in header):
                print(f"   ⚠️ 表格0不是批量汇总表")
                return

        # 提取数据行
        for row_idx, row in enumerate(table.rows[1:], 1):
            cells = [c.text.strip() for c in row.cells]

            if len(cells) < 4:
                continue

            # 跳过空行和合计行
            if not cells[0] or not cells[0].isdigit():
                continue

            subject = BatchSubject()

            try:
                subject.seq_no = int(cells[0])
            except:
                continue

            subject.address = cells[1]

            try:
                subject.building_area = float(cells[2])
            except:
                pass

            try:
                subject.total_price = float(cells[3])
            except:
                pass

            # 计算单价
            if subject.building_area > 0 and subject.total_price > 0:
                subject.unit_price = subject.total_price * 10000 / subject.building_area

            result.subjects.append(subject)

        result.total_count = len(result.subjects)

    def _extract_case_groups(self, result: XianzhibExtractionResult):
        """提取可比实例组"""
        # 查找包含"可比实例A/B/C"表头的表格
        for t_idx, table in enumerate(self.tables):
            if len(table.rows) == 0:
                continue

            header = ' '.join([c.text.strip() for c in table.rows[0].cells[:6]])

            # 检查是否是基础信息表（包含项目、估价对象、可比实例）
            if '项目' in header and '估价对象' in header and '可比实例' in header:
                cases = self._extract_case_group(table, t_idx)
                if cases:
                    result.case_groups.append(cases)

    def _extract_case_group(self, table, table_idx: int) -> List[Case]:
        """提取单组可比实例"""
        cases = [Case(case_id='A'), Case(case_id='B'), Case(case_id='C')]

        COL_A = 3
        COL_B = 4
        COL_C = 5

        for row_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]

            if len(cells) < 6:
                continue

            label = cells[0] + cells[1] if len(cells) > 1 else cells[0]
            label = label.replace(' ', '').replace('\u3000', '')

            if '地址' in label or '坐落' in label:
                for i, case in enumerate(cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.address = LocatedValue(
                            value=cells[col],
                            position=Position(table_idx, row_idx, col),
                            raw_text=cells[col]
                        )

            elif '成交基价' in label or '交易价格' in label:
                for i, case in enumerate(cases):
                    col = COL_A + i
                    if col < len(cells):
                        try:
                            price = float(re.sub(r'[^\d.]', '', cells[col]))
                            case.transaction_price = LocatedValue(
                                value=price,
                                position=Position(table_idx, row_idx, col),
                                raw_text=cells[col]
                            )
                        except:
                            pass

            elif '建筑面积' in label:
                for i, case in enumerate(cases):
                    col = COL_A + i
                    if col < len(cells):
                        try:
                            area = float(re.sub(r'[^\d.]', '', cells[col]))
                            case.building_area = LocatedValue(
                                value=area,
                                position=Position(table_idx, row_idx, col),
                                raw_text=cells[col]
                            )
                        except:
                            pass

            elif '交易日期' in label:
                for i, case in enumerate(cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.transaction_date = cells[col]

            elif '来源' in label:
                for i, case in enumerate(cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.data_source = cells[col]

        return cases

    def _extract_floor_corrections(self, result: XianzhibExtractionResult):
        """提取楼层修正信息"""
        # 查找楼层修正结果表
        for t_idx, table in enumerate(self.tables):
            if len(table.rows) == 0:
                continue

            header = ' '.join([c.text.strip() for c in table.rows[0].cells[:7]])

            # 楼层修正表应该包含：坐落、建筑面积、基准价、楼层系数、单价、总价
            if '基准价' in header and '楼层系数' in header:
                self._parse_floor_correction_table(table, result)
                break

    def _parse_floor_correction_table(self, table, result: XianzhibExtractionResult):
        """解析楼层修正表"""
        for row_idx, row in enumerate(table.rows[1:], 1):
            cells = [c.text.strip() for c in row.cells]

            if len(cells) < 6:
                continue

            address = cells[0]

            # 查找对应的估价对象
            for subject in result.subjects:
                if address in subject.address or subject.address in address:
                    # 提取楼层系数
                    try:
                        # 楼层系数通常在第4或5列
                        for cell in cells[3:6]:
                            if '%' in cell or (cell.replace('.', '').isdigit() and float(cell) < 2):
                                factor = float(cell.replace('%', ''))
                                if factor > 10:  # 百分比形式
                                    factor = factor / 100
                                subject.floor_factor = factor
                                break
                    except:
                        pass
                    break

    def _normalize_factor_name(self, name: str) -> str:
        """标准化因素名称"""
        name = name.replace(' ', '').replace('\u3000', '')

        mapping = {
            '区域位置': 'location_region',
            '楼幢位置': 'location_building',
            '朝向': 'orientation',
            '交通条件': 'traffic',
            '配套设施': 'facilities',
            '环境质量': 'environment',
            '景观': 'landscape',
            '物业管理': 'property_management',
            '建筑面积': 'area',
            '空间布局': 'layout',
            '新旧程度': 'age',
            '装饰装修': 'decoration',
            '建筑结构': 'structure',
            '物业类型': 'property_type',
            '设施设备': 'equipment',
        }

        return mapping.get(name, name)


# ============================================================================
# 测试
# ============================================================================

if __name__ == "__main__":
    extractor = XianzhibExtractor()
    result = extractor.extract("./data/docs/批量评估报告.docx")

    print(f"\n{'=' * 70}")
    print("【提取结果】")
    print('=' * 70)

    print(f"\n估价对象数量: {result.total_count}")
    print(f"总建筑面积: {result.total_area:.2f}㎡")
    print(f"总评估价值: {result.total_value:.2f}万元")

    print(f"\n估价对象列表:")
    for subject in result.subjects[:5]:
        print(f"  {subject.seq_no}. {subject.address}")
        print(f"     面积: {subject.building_area}㎡, 总价: {subject.total_price}万元")
        if subject.floor_factor != 1.0:
            print(f"     楼层系数: {subject.floor_factor}")

    if len(result.subjects) > 5:
        print(f"  ... 还有 {len(result.subjects) - 5} 个")

    print(f"\n可比实例组: {len(result.case_groups)}组")
    for i, cases in enumerate(result.case_groups):
        print(f"  第{i + 1}组:")
        for case in cases:
            print(f"    {case.case_id}: {case.address.value}, 成交价: {case.transaction_price.value}元/㎡")

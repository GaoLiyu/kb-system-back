"""
租金报告精确提取器
==================
针对租金报告的表格结构精确提取
表格索引（基于分析结果）：
- 表格0/2/9: 结果汇总表（坐落、评估面积、评估单价、评估总价）
- 表格1: 权属表
- 表格4: 基础信息表
- 表格5: 因素描述表
- 表格6: 因素等级表
- 表格7: 因素指数表（注意：表头是"案例A/B/C"而不是"可比实例"）
- 表格8: 修正计算表
"""

import os
import re
from typing import Dict, List
from docx import Document
from dataclasses import dataclass, field

from .table_utils import extract_property_rights_generic


@dataclass
class Position:
    """位置信息"""
    table_index: int = -1
    row_index: int = -1
    col_index: int = -1


@dataclass
class LocatedValue:
    """带位置的值"""
    value: any = None
    position: Position = field(default_factory=Position)
    raw_text: str = ""


@dataclass
class Factor:
    """因素数据"""
    name: str = ""
    description: str = ""
    level: str = ""
    index: int = 100
    desc_pos: Position = field(default_factory=Position)
    level_pos: Position = field(default_factory=Position)
    index_pos: Position = field(default_factory=Position)


@dataclass
class Case:
    """可比实例"""
    case_id: str = ""
    address: LocatedValue = field(default_factory=LocatedValue)
    location: str = ""
    usage: str = ""
    data_source: str = ""
    rental_price: LocatedValue = field(default_factory=LocatedValue)  # 租赁价格
    building_area: LocatedValue = field(default_factory=LocatedValue)
    transaction_date: str = ""

    # 新增字段
    district: str = ""           # 区域
    street: str = ""             # 街道/镇
    build_year: int = 0          # 建成年份
    total_floor: int = 0         # 总楼层
    current_floor: str = ""       # 所在楼层
    structure: str = ""          # 结构
    orientation: str = ""        # 朝向
    decoration: str = ""         # 装修

    # 新增价格内涵里面的字段
    property_scope: str = ""  # 财产范围
    payment_methods: str = ""  # 付款方式
    financing_conditions: str = ""  # 融资条件
    tax_burden: str = ""  # 税负
    unit_measurement: str = ""  # 计价单位
    price_type: str = ""  # 价格类型

    # 修正系数
    transaction_correction: LocatedValue = field(default_factory=LocatedValue)
    market_correction: LocatedValue = field(default_factory=LocatedValue)
    location_correction: LocatedValue = field(default_factory=LocatedValue)
    physical_correction: LocatedValue = field(default_factory=LocatedValue)
    rights_correction: LocatedValue = field(default_factory=LocatedValue)
    adjusted_price: LocatedValue = field(default_factory=LocatedValue)

    # 因素
    location_factors: Dict[str, Factor] = field(default_factory=dict)
    physical_factors: Dict[str, Factor] = field(default_factory=dict)
    rights_factors: Dict[str, Factor] = field(default_factory=dict)


@dataclass
class Subject:
    """估价对象"""
    address: LocatedValue = field(default_factory=LocatedValue)
    building_area: LocatedValue = field(default_factory=LocatedValue)
    unit_price: LocatedValue = field(default_factory=LocatedValue)   # 元/㎡·年
    total_price: LocatedValue = field(default_factory=LocatedValue)  # 万元/年

    # 权属
    cert_no: str = "" # 房屋所有权证证号
    owner: str = "" # 房屋所有权人
    structure: str = "" # 结构
    floor: str = "" # 楼层
    plan_usage: str = "" # 规划用途
    land_no: str = "" # 土地使用证证号
    land_owner: str = "" # 土地使用权人
    land_address: str = "" # 土地坐落
    land_use_type: str = "" # 使用权类型
    land_type: str = "" # 地类（用途）
    land_area: float = 0.0 # 土地面积
    end_date: str = "" # 终止日期

    # 新增字段
    district: str = ""            # 区域
    street: str = ""              # 街道/镇
    location: str = ""            # 位置
    data_source: str = ""         # 数据来源
    build_year: int = 0           # 建成年份
    total_floor: str = ""         # 总楼层（字符串）
    current_floor: str = ""       # 所在楼层（字符串）
    orientation: str = ""         # 朝向
    decoration: str = ""          # 装修
    value_date: str = ""          # 价值时点
    appraisal_purpose: str = ""   # 估价目的
    usage: str = ""               # 用途

    # 新增价格内涵里面的字段
    property_scope: str = "" # 财产范围
    payment_methods: str = "" # 付款方式
    financing_conditions: str = "" # 融资条件
    tax_burden: str = "" # 税负
    unit_measurement: str = "" # 计价单位
    price_type: str = "" # 价格类型
    rental_price: LocatedValue = field(default_factory=LocatedValue) # 租赁价格
    transaction_date: str = ""

    # 因素
    location_factors: Dict[str, Factor] = field(default_factory=dict)
    physical_factors: Dict[str, Factor] = field(default_factory=dict)
    rights_factors: Dict[str, Factor] = field(default_factory=dict)


@dataclass
class ZujinExtractionResult:
    """租金报告提取结果"""
    source_file: str = ""
    subject: Subject = field(default_factory=Subject)
    cases: List[Case] = field(default_factory=list)

    final_unit_price: LocatedValue = field(default_factory=LocatedValue)
    final_total_price: LocatedValue = field(default_factory=LocatedValue)
    price_unit: str = "元/㎡·年"

    # ✅ 你代码里用到了 result.floor_factor，所以这里补上
    floor_factor: float = 1.0

    type: str = "" # 类型


class ZujinExtractor:
    """租金报告提取器"""

    # 表格索引
    TABLE_RESULT_SUMMARY = 0
    TABLE_PROPERTY_RIGHTS = 1
    TABLE_BASIC_INFO = 4
    TABLE_FACTOR_DESC = 5
    TABLE_FACTOR_LEVEL = 6
    TABLE_FACTOR_INDEX = 7
    TABLE_CORRECTION = 8

    # 因素名称（租金报告特有的因素）
    LOCATION_FACTORS = ['繁华程度', '区域位置', '楼幢位置', '朝向', '交通条件', '配套设施', '环境质量', '景观',
                        '物业管理', '驻车条件']
    PHYSICAL_FACTORS = ['建筑面积', '套内建筑面积', '空间布局', '新旧程度', '装饰装修', '建筑结构', '建筑结构',
                        '物业类型', '设施设备', '楼宇等级', '地形地势', '地质土壤', '开发程度']
    RIGHTS_FACTORS = ['规划条件', '土地使用期限', '土地剩余使用年限', '担保物权设立', '租赁占用状况', '拖欠税费状况',
                      '登记状况', '他项权利', '限制权利', '其他因素']

    def __init__(self, auto_detect: bool = True):
        self.doc = None
        self.tables = []
        self.full_text = ""
        self.auto_detect = auto_detect

    def extract(self, doc_path: str) -> ZujinExtractionResult:
        """提取租金报告"""
        self.doc = Document(doc_path)
        self.tables = self.doc.tables
        self.full_text = "\n".join([p.text for p in self.doc.paragraphs])

        result = ZujinExtractionResult(source_file=os.path.basename(doc_path), type='zujin')

        print(f"\n📊 提取租金报告: {os.path.basename(doc_path)}")
        print(f"   表格数量: {len(self.tables)}")

        # 自动检测表格索引
        if self.auto_detect:
            self._auto_detect_table_indices()
            print(
                f"   ✓ 自动检测表格索引: 汇总={self.TABLE_RESULT_SUMMARY}, 权属={self.TABLE_PROPERTY_RIGHTS}, "
                f"基础={self.TABLE_BASIC_INFO}, 描述={self.TABLE_FACTOR_DESC}, 等级={self.TABLE_FACTOR_LEVEL}, "
                f"指数={self.TABLE_FACTOR_INDEX}, 修正={self.TABLE_CORRECTION}"
            )

        # 1. 提取结果汇总
        self._extract_result_summary(result)
        print(f"   ✓ 结果汇总: {result.subject.address.value}")

        # 2. 提取权属信息（表1）
        self._extract_property_rights(result)

        # 3. 提取基础信息
        self._extract_basic_info(result)
        print(f"   ✓ 基础信息: {len(result.cases)}个可比实例")

        # 4. 提取因素描述
        self._extract_factor_descriptions(result)

        # 5. 提取因素等级
        self._extract_factor_levels(result)

        # 6. 提取因素指数
        self._extract_factor_indices(result)
        print(f"   ✓ 因素数据: 描述/等级/指数")

        # 7. 提取修正系数
        self._extract_corrections(result)
        print(f"   ✓ 修正系数")

        # 8. 提取楼层修正系数（从全文匹配，和 shezhi 一样）
        self._extract_floor_factor(result)
        if result.floor_factor != 1.0:
            print(f"   ✓ 楼层修正: {result.floor_factor}")

        # 9. 扩展信息（价值时点/估价目的/建成年代等）
        self._extract_extended_info(result)

        # 10. 解析区域信息
        self._parse_district(result)

        # 11. 补一次楼层解析（如果只拿到了 floor 原始串）
        self._parse_floor_from_floor_str(result.subject)

        return result

    # ----------------- 工具函数（和 shezhi 同样风格） -----------------

    @staticmethod
    def _norm_num_str(s: str) -> str:
        return re.sub(r"[^\d]", "", s or "")

    def _auto_detect_table_indices(self):
        """自动检测租金报告关键表格索引（打分制，抗位置偏移）"""

        def norm(s: str) -> str:
            if not s:
                return ""
            return (
                s.replace("\u3000", " ")
                .replace("\n", " ")
                .replace("\t", " ")
                .strip()
            )

        def table_block(table, max_rows=10, max_cols=12) -> str:
            parts = []
            rN = min(len(table.rows), max_rows)
            for r in range(rN):
                row = table.rows[r]
                cN = min(len(row.cells), max_cols)
                for c in range(cN):
                    parts.append(norm(row.cells[c].text))
            return " ".join([p for p in parts if p])

        def compact(s: str) -> str:
            return s.replace(" ", "")

        def count_hits(text: str, keys) -> int:
            return sum(1 for k in keys if k in text)

        def has_all(text: str, keys) -> bool:
            return all(k in text for k in keys)

        best = {
            "summary": (-1, self.TABLE_RESULT_SUMMARY),
            "rights": (-1, self.TABLE_PROPERTY_RIGHTS),
            "basic": (-1, self.TABLE_BASIC_INFO),
            "desc": (-1, self.TABLE_FACTOR_DESC),
            "level": (-1, self.TABLE_FACTOR_LEVEL),
            "index": (-1, self.TABLE_FACTOR_INDEX),
            "corr": (-1, self.TABLE_CORRECTION),
        }

        for i, table in enumerate(self.tables):
            if len(table.rows) == 0:
                continue

            rows = len(table.rows)
            cols = len(table.columns) if table.columns else 0

            block = table_block(table, max_rows=12, max_cols=12)
            t = compact(block)

            # ---------------- 1) 结果汇总表 ----------------
            # 强特征：坐落 + 评估面积 + 评估单价 + 评估总价 + 元/㎡·年（或 万元/年）
            score_summary = 0
            if has_all(t, ["评估面积", "评估单价", "评估总价"]):
                score_summary += 10
            score_summary += count_hits(t, ["坐落", "元/㎡·年", "万元/年", "（㎡）"])
            if rows <= 6:
                score_summary += 2
            if score_summary > best["summary"][0]:
                best["summary"] = (score_summary, i)

            # ---------------- 2) 权属表 ----------------
            # 典型字段：房屋所有权证证号/房屋所有权人/土地使用证证号/土地使用权人/终止日期 等
            score_rights = 0
            score_rights += 3 * count_hits(t, ["房屋所有权证证号", "房屋所有权人"])
            score_rights += 3 * count_hits(t, ["土地使用证证号", "土地使用权人"])
            score_rights += count_hits(t, ["规划用途", "建筑面积", "坐落", "使用权类型", "地类", "土地使用权面积",
                                           "终止日期"])
            if score_rights > best["rights"][0]:
                best["rights"] = (score_rights, i)

            # ---------------- 3) 基础信息表 ----------------
            # 租金模板通常包含：项目/估价对象/可比实例 或 案例A/B/C 结构
            score_basic = 0
            if ("项目" in t and "估价对象" in t and ("可比实例" in t or "案例A" in t)):
                score_basic += 10
            score_basic += count_hits(t, ["地址", "位置", "来源", "用途", "财产范围", "付款方式", "融资条件", "税负",
                                          "计价单位", "价格类型", "租赁价格", "交易日期"])
            if rows >= 10 and cols >= 5:
                score_basic += 2
            if score_basic > best["basic"][0]:
                best["basic"] = (score_basic, i)

            # ---------------- 4) 因素表：描述/等级/指数 ----------------
            # 共同特征：区位状况/实物状况/权益状况 + A/B/C列头（可比实例A/B/C 或 案例A/B/C）
            has_abc = (("可比实例A" in t and "可比实例B" in t and "可比实例C" in t) or
                       ("案例A" in t and "案例B" in t and "案例C" in t))
            has_factors = ("区位状况" in t or "实物状况" in t or "权益状况" in t)

            if has_abc and has_factors:
                # 4.1 描述表：更多“交通/配套/环境/装修/结构/新旧”等描述词
                score_desc = 0
                score_desc += count_hits(t, ["交通条件", "配套设施", "环境质量", "物业管理", "驻车条件", "装饰装修",
                                             "建筑结构", "新旧程度", "空间布局"])
                # 描述表通常不强调“指数/100”
                score_desc -= 2 * count_hits(t, ["指数", "100"])
                if score_desc > best["desc"][0]:
                    best["desc"] = (score_desc, i)

                # 4.2 等级表：等级/优良中差 等
                score_level = 0
                score_level += 3 * count_hits(t, ["等级"])
                score_level += count_hits(t, ["优", "良", "中", "差", "较优", "一般"])
                score_level -= count_hits(t, ["指数", "100"])
                if score_level > best["level"][0]:
                    best["level"] = (score_level, i)

                # 4.3 指数表：指数/100 特别多
                score_index = 0
                score_index += 4 * count_hits(t, ["指数"])
                if "100" in t:
                    score_index += 2
                if score_index > best["index"][0]:
                    best["index"] = (score_index, i)

            # ---------------- 5) 修正计算表 ----------------
            # 强特征：比较因素修正表 + 交易价格/交易情况修正系数/市场状况调整系数/调整后单价（元/㎡·年）
            score_corr = 0
            if "比较因素修正表" in t:
                score_corr += 10
            score_corr += 2 * count_hits(t, ["交易价格", "交易情况修正系数", "市场状况调整系数", "区位状况调整系数",
                                             "实物状况调整系数", "权益状况调整系数"])
            score_corr += count_hits(t, ["调整后单价", "元/㎡·年"])
            if rows >= 6 and cols >= 4:
                score_corr += 2
            if score_corr > best["corr"][0]:
                best["corr"] = (score_corr, i)

        # --------- 落盘（阈值防误判）---------
        if best["summary"][0] >= 10:
            self.TABLE_RESULT_SUMMARY = best["summary"][1]
        if best["rights"][0] >= 6:
            self.TABLE_PROPERTY_RIGHTS = best["rights"][1]
        if best["basic"][0] >= 10:
            self.TABLE_BASIC_INFO = best["basic"][1]
        if best["desc"][0] >= 2:
            self.TABLE_FACTOR_DESC = best["desc"][1]
        if best["level"][0] >= 2:
            self.TABLE_FACTOR_LEVEL = best["level"][1]
        if best["index"][0] >= 2:
            self.TABLE_FACTOR_INDEX = best["index"][1]
        if best["corr"][0] >= 10:
            self.TABLE_CORRECTION = best["corr"][1]

        # 兜底：如果因素表没识别全，按“基础信息表后续顺序”兜一下（但不强依赖）
        base = self.TABLE_BASIC_INFO
        if 0 <= base < len(self.tables):
            if best["desc"][0] < 2:
                self.TABLE_FACTOR_DESC = min(base + 1, len(self.tables) - 1)
            if best["level"][0] < 2:
                self.TABLE_FACTOR_LEVEL = min(base + 2, len(self.tables) - 1)
            if best["index"][0] < 2:
                self.TABLE_FACTOR_INDEX = min(base + 3, len(self.tables) - 1)
            if best["corr"][0] < 10:
                self.TABLE_CORRECTION = min(base + 4, len(self.tables) - 1)

    def _set_subject_floor(self, subject: Subject, cur: str, total: str):
        # cur_n = self._norm_num_str(cur)
        # total_n = self._norm_num_str(total)
        if cur:
            subject.current_floor = cur
        if total:
            subject.total_floor = total

    def _parse_floor_from_floor_str(self, subject: Subject):
        """从 subject.floor 解析 current/total（字符串）"""
        if not subject.floor:
            return
        text = subject.floor.strip()
        # m = re.search(r"(\d+)(?:-(\d+))?\s*/\s*(\d+)", text)
        if text:
            # cur = m.group(2) or m.group(1)
            # total = m.group(3)
            cur = text.split("/")[0]
            total = text.split("/")[1]
            self._set_subject_floor(subject, cur, total)

    # ----------------- 表格提取 -----------------

    def _extract_result_summary(self, result: ZujinExtractionResult):
        """提取结果汇总表（支持有/无楼层列）"""
        table = self.tables[self.TABLE_RESULT_SUMMARY]
        if len(table.rows) < 2:
            return

        header = [c.text.strip() for c in table.rows[0].cells]
        has_floor_col = any('楼层' in h for h in header)

        row = table.rows[1]
        cells = [c.text.strip() for c in row.cells]

        if has_floor_col:
            col_address, col_floor, col_area, col_unit_price, col_total_price = 0, 1, 2, 3, 4
        else:
            col_address, col_floor, col_area, col_unit_price, col_total_price = 0, -1, 1, 2, 3

        # 地址
        if len(cells) > col_address:
            result.subject.address = LocatedValue(
                value=cells[col_address],
                position=Position(self.TABLE_RESULT_SUMMARY, 1, col_address),
                raw_text=cells[col_address]
            )

        # 楼层（如果存在）
        if has_floor_col and col_floor >= 0 and len(cells) > col_floor:
            floor_text = cells[col_floor]
            if floor_text:
                result.subject.floor = floor_text
                self._parse_floor_from_floor_str(result.subject)

        # 面积
        if len(cells) > col_area:
            try:
                result.subject.building_area = LocatedValue(
                    value=float(re.sub(r"[^\d.]", "", cells[col_area])),
                    position=Position(self.TABLE_RESULT_SUMMARY, 1, col_area),
                    raw_text=cells[col_area]
                )
            except:
                pass

        # 单价
        if len(cells) > col_unit_price:
            try:
                result.subject.unit_price = LocatedValue(
                    value=float(re.sub(r"[^\d.]", "", cells[col_unit_price])),
                    position=Position(self.TABLE_RESULT_SUMMARY, 1, col_unit_price),
                    raw_text=cells[col_unit_price]
                )
                result.final_unit_price = result.subject.unit_price
            except:
                pass

        # 总价
        if len(cells) > col_total_price:
            try:
                result.subject.total_price = LocatedValue(
                    value=float(re.sub(r"[^\d.]", "", cells[col_total_price])),
                    position=Position(self.TABLE_RESULT_SUMMARY, 1, col_total_price),
                    raw_text=cells[col_total_price]
                )
                result.final_total_price = result.subject.total_price
            except:
                pass

    def _extract_property_rights(self, result: ZujinExtractionResult):
        """提取权属表（使用 table_utils：表头定位 + 列映射，不改变你的结果结构）"""
        if len(self.tables) <= self.TABLE_PROPERTY_RIGHTS:
            return

        table = self.tables[self.TABLE_PROPERTY_RIGHTS]

        def subject_setter(key: str, value):
            # ---- 房屋块 ----
            if key == "cert_no":
                if value and not result.subject.cert_no:
                    result.subject.cert_no = str(value).strip()

            elif key == "owner":
                if value and not result.subject.owner:
                    result.subject.owner = str(value).strip()

            elif key == "address":
                # zujin 的 Subject.address 也是 LocatedValue
                if value and not result.subject.address.value:
                    result.subject.address.value = str(value).strip()

            elif key == "structure":
                if value and not result.subject.structure:
                    result.subject.structure = str(value).strip()

            elif key == "floor":
                if value and not result.subject.floor:
                    result.subject.floor = str(value).strip()
                    # 你原来会解析 current/total，这里也保持原行为
                    self._parse_floor_from_floor_str(result.subject)

            elif key == "plan_usage":
                if value and not result.subject.plan_usage:
                    result.subject.plan_usage = str(value).strip()

            elif key == "building_area":
                # 你原来写的是 result.subject.building_area (LocatedValue)
                if value is not None and not result.subject.building_area.value:
                    result.subject.building_area.value = float(value)

            # ---- 土地块（zujin Subject 里字段更全）----
            elif key == "land_no":
                if value and not result.subject.land_no:
                    result.subject.land_no = str(value).strip()

            elif key == "land_owner":
                if value and not result.subject.land_owner:
                    result.subject.land_owner = str(value).strip()

            elif key == "land_address":
                if value and not result.subject.land_address:
                    result.subject.land_address = str(value).strip()

            elif key == "land_use_type":
                if value and not result.subject.land_use_type:
                    result.subject.land_use_type = str(value).strip()

            elif key == "land_type":
                if value and not result.subject.land_type:
                    result.subject.land_type = str(value).strip()

            elif key == "land_area":
                if value is not None and not result.subject.land_area:
                    try:
                        result.subject.land_area = float(value)
                    except:
                        pass

            elif key == "end_date":
                if value and not result.subject.end_date:
                    result.subject.end_date = str(value).strip()

        extract_property_rights_generic(table, subject_setter=subject_setter, detect_land=True)

    def _extract_basic_info(self, result: ZujinExtractionResult):
        """提取基础信息表"""
        table = self.tables[self.TABLE_BASIC_INFO]
        result.cases = [Case(case_id='A'), Case(case_id='B'), Case(case_id='C')]

        COL_SUBJECT, COL_A, COL_B, COL_C = 2, 3, 4, 5

        for row_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 6:
                continue

            label = (cells[0] + cells[1]).replace(' ', '').replace('\u3000', '')

            if '地址' in label or '坐落' in label:
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.address = LocatedValue(
                            value=cells[col],
                            position=Position(self.TABLE_BASIC_INFO, row_idx, col),
                            raw_text=cells[col]
                        )

            elif '位置' in label and '楼' not in label:
                if COL_SUBJECT < len(cells):
                    result.subject.location = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.location = cells[col]

            elif '来源' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.data_source = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.data_source = cells[col]

            elif '用途' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.usage = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.usage = cells[col]

            elif '财产范围' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.property_scope = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.property_scope = cells[col]

            elif '付款方式' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.payment_methods = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.payment_methods = cells[col]

            elif '融资条件' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.financing_conditions = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.financing_conditions = cells[col]

            elif '税' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.tax_burden = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.tax_burden = cells[col]

            elif '单位' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.unit_measurement = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.unit_measurement = cells[col]

            elif '价格类型' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.price_type = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.price_type = cells[col]

            elif '租赁价格' in label or '交易价格' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.price = LocatedValue(
                        value=cells[COL_SUBJECT],
                        position=Position(self.TABLE_BASIC_INFO, row_idx, COL_SUBJECT),
                        raw_text=cells[COL_SUBJECT]
                    )
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        try:
                            price = float(re.sub(r'[^\d.]', '', cells[col]))
                            case.rental_price = LocatedValue(
                                value=price,
                                position=Position(self.TABLE_BASIC_INFO, row_idx, col),
                                raw_text=cells[col]
                            )
                        except:
                            pass

            elif '评估面积' in label or '建筑面积' in label:
                if COL_SUBJECT < len(cells):
                    result.subject.building_area = LocatedValue(
                        value=cells[COL_SUBJECT],
                        position=Position(self.TABLE_BASIC_INFO, row_idx, COL_SUBJECT),
                        raw_text=cells[COL_SUBJECT]
                    )
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        try:
                            area = float(re.sub(r'[^\d.]', '', cells[col]))
                            case.building_area = LocatedValue(
                                value=area,
                                position=Position(self.TABLE_BASIC_INFO, row_idx, col),
                                raw_text=cells[col]
                            )
                        except:
                            pass

            elif '交易日期' in label or '成交日期' in label or '交易时间' in label or '价值时点' in label:
                result.subject.transaction_date = cells[COL_SUBJECT]
                for i, case in enumerate(result.cases):
                    col = COL_A + i
                    if col < len(cells):
                        case.transaction_date = cells[col]

    def _extract_factor_descriptions(self, result: ZujinExtractionResult):
        """提取因素描述表（表5）——按固定6列读取"""
        table = self.tables[self.TABLE_FACTOR_DESC]

        COL_CATEGORY, COL_FACTOR, COL_SUBJECT, COL_A, COL_B, COL_C = 0, 1, 2, 3, 4, 5

        category_alias = {
            '区位状况': '区位状况',
            '实物状况': '实物状况',
            '实物因素': '实物状况',
            '权益状况': '权益状况',
            '权益因素': '权益状况',
        }

        current_category = ''
        for row_idx, row in enumerate(table.rows[1:], 1):
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            if len(cells) < 6:
                continue

            raw_category = (cells[COL_CATEGORY] or '').replace(' ', '').replace('　', '')
            factor_name = (cells[COL_FACTOR] or '').replace(' ', '').replace('　', '')

            if raw_category in ('交易情况', '交易日期') or factor_name in ('交易情况', '交易日期'):
                continue

            if raw_category in category_alias:
                current_category = category_alias[raw_category]

            if not factor_name:
                continue

            factor_type = self._get_factor_type(factor_name, current_category)
            if not factor_type:
                continue

            factor_key = self._normalize_factor_name(factor_name)

            # 估价对象
            subject_val = cells[COL_SUBJECT]
            if subject_val:
                subject_dict = getattr(result.subject, f'{factor_type}_factors')
                f = subject_dict.get(factor_key) or Factor(name=factor_key)
                f.description = subject_val
                f.desc_pos = Position(self.TABLE_FACTOR_DESC, row_idx, COL_SUBJECT)
                subject_dict[factor_key] = f

                # ✅ 关键修复：传 subject_val（不是 factor_type）
                self._sync_subject_fields_from_factor(result.subject, factor_key, subject_val)

            # 可比实例A/B/C
            for i, case in enumerate(result.cases):
                col = [COL_A, COL_B, COL_C][i]
                value = cells[col]
                if value == '':
                    continue
                factor_dict = getattr(case, f'{factor_type}_factors')
                f = factor_dict.get(factor_key) or Factor(name=factor_key)
                f.description = value
                f.desc_pos = Position(self.TABLE_FACTOR_DESC, row_idx, col)
                factor_dict[factor_key] = f

                # ✅ 关键修复：传 value（不是 factor_type）
                self._sync_case_fields_from_factor(case, factor_key, value)

    def _extract_factor_levels(self, result: ZujinExtractionResult):
        """提取因素等级表（表6）——按固定6列读取"""
        table = self.tables[self.TABLE_FACTOR_LEVEL]
        COL_CATEGORY, COL_FACTOR, COL_SUBJECT, COL_A, COL_B, COL_C = 0, 1, 2, 3, 4, 5

        category_alias = {
            '区位状况': '区位状况',
            '实物状况': '实物状况',
            '实物因素': '实物状况',
            '权益状况': '权益状况',
            '权益因素': '权益状况',
        }

        current_category = ''
        for row_idx, row in enumerate(table.rows[1:], 1):
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            if len(cells) < 6:
                continue

            raw_category = (cells[COL_CATEGORY] or '').replace(' ', '').replace('　', '')
            factor_name = (cells[COL_FACTOR] or '').replace(' ', '').replace('　', '')

            if raw_category in ('交易情况', '交易日期') or factor_name in ('交易情况', '交易日期'):
                continue

            if raw_category in category_alias:
                current_category = category_alias[raw_category]

            if not factor_name:
                continue

            factor_type = self._get_factor_type(factor_name, current_category)
            if not factor_type:
                continue

            factor_key = self._normalize_factor_name(factor_name)

            # 估价对象
            subject_val = cells[COL_SUBJECT]
            if subject_val:
                subject_dict = getattr(result.subject, f'{factor_type}_factors')
                f = subject_dict.get(factor_key) or Factor(name=factor_key)
                f.level = subject_val
                f.level_pos = Position(self.TABLE_FACTOR_LEVEL, row_idx, COL_SUBJECT)
                subject_dict[factor_key] = f

            # 可比实例
            for i, case in enumerate(result.cases):
                col = [COL_A, COL_B, COL_C][i]
                value = cells[col]
                if value == '':
                    continue
                factor_dict = getattr(case, f'{factor_type}_factors')
                f = factor_dict.get(factor_key) or Factor(name=factor_key)
                f.level = value
                f.level_pos = Position(self.TABLE_FACTOR_LEVEL, row_idx, col)
                factor_dict[factor_key] = f

    def _extract_factor_indices(self, result: ZujinExtractionResult):
        """提取因素指数表（表7）——按固定6列读取"""
        table = self.tables[self.TABLE_FACTOR_INDEX]
        COL_CATEGORY, COL_FACTOR, COL_SUBJECT, COL_A, COL_B, COL_C = 0, 1, 2, 3, 4, 5

        category_alias = {
            '区位状况': '区位状况',
            '实物状况': '实物状况',
            '实物因素': '实物状况',
            '权益状况': '权益状况',
            '权益因素': '权益状况',
        }

        def to_int(v: str) -> int:
            try:
                return int(re.sub(r'[^0-9]', '', v))
            except Exception:
                return 100

        current_category = ''
        for row_idx, row in enumerate(table.rows[1:], 1):
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            if len(cells) < 6:
                continue

            raw_category = (cells[COL_CATEGORY] or '').replace(' ', '').replace('　', '')
            factor_name = (cells[COL_FACTOR] or '').replace(' ', '').replace('　', '')

            if raw_category in ('交易情况', '交易日期') or factor_name in ('交易情况', '交易日期'):
                continue

            if raw_category in category_alias:
                current_category = category_alias[raw_category]

            if not factor_name:
                continue

            factor_type = self._get_factor_type(factor_name, current_category)
            if not factor_type:
                continue

            factor_key = self._normalize_factor_name(factor_name)

            # 估价对象
            subject_val = cells[COL_SUBJECT]
            if subject_val:
                subject_dict = getattr(result.subject, f'{factor_type}_factors')
                f = subject_dict.get(factor_key) or Factor(name=factor_key)
                f.index = to_int(subject_val)
                f.index_pos = Position(self.TABLE_FACTOR_INDEX, row_idx, COL_SUBJECT)
                subject_dict[factor_key] = f

            # 可比实例
            for i, case in enumerate(result.cases):
                col = [COL_A, COL_B, COL_C][i]
                value = cells[col]
                if value == '':
                    continue
                factor_dict = getattr(case, f'{factor_type}_factors')
                f = factor_dict.get(factor_key) or Factor(name=factor_key)
                f.index = to_int(value)
                f.index_pos = Position(self.TABLE_FACTOR_INDEX, row_idx, col)
                factor_dict[factor_key] = f

    def _extract_corrections(self, result: ZujinExtractionResult):
        """提取修正系数"""
        table = self.tables[self.TABLE_CORRECTION]
        COL_A = 1

        ROW_MAPPING = {
            '交易价格': 'rental_price',
            '交易情况修正': 'transaction_correction',
            '市场状况': 'market_correction',
            '区位状况': 'location_correction',
            '实物状况': 'physical_correction',
            '权益状况': 'rights_correction',
            '调整后单价': 'adjusted_price',
        }

        for row_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2:
                continue

            label = cells[0].replace(' ', '').replace('\u3000', '')
            field_name = None
            for key, field in ROW_MAPPING.items():
                if key in label:
                    field_name = field
                    break
            if not field_name:
                continue

            for i, case in enumerate(result.cases):
                col = COL_A + i
                if col < len(cells):
                    try:
                        value = float(re.sub(r"[^\d.]", "", cells[col]))
                        loc_val = LocatedValue(
                            value=value,
                            position=Position(self.TABLE_CORRECTION, row_idx, col),
                            raw_text=cells[col]
                        )
                        setattr(case, field_name, loc_val)
                    except:
                        pass

    # ----------------- 规则/同步（保持你 shezhi 的写法） -----------------

    def _get_factor_type(self, factor_name: str, current_category: str) -> str:
        """获取因素类型"""
        if factor_name in self.LOCATION_FACTORS or current_category == '区位状况':
            return 'location'
        elif factor_name in self.PHYSICAL_FACTORS or current_category in ('实物状况', '实物因素'):
            return 'physical'
        elif factor_name in self.RIGHTS_FACTORS or current_category in ('权益状况', '权益因素'):
            return 'rights'
        return ''

    def _sync_subject_fields_from_factor(self, subject: Subject, factor_key: str, val: str):
        """把因素描述同步到 Subject 的新增字段（不改变输出结构）"""
        v = (val or "").strip()
        if not v:
            return
        if factor_key == "orientation" and not subject.orientation:
            subject.orientation = v
        elif factor_key == "decoration" and not subject.decoration:
            subject.decoration = v
        elif factor_key == "structure" and not subject.structure:
            subject.structure = v

    def _sync_case_fields_from_factor(self, case: Case, factor_key: str, val: str):
        """把因素描述同步到 Case 的新增字段（不改变输出结构）"""
        v = (val or "").strip()
        if not v:
            return
        if factor_key == "orientation" and not case.orientation:
            case.orientation = v
        elif factor_key == "decoration" and not case.decoration:
            case.decoration = v
        elif factor_key == "structure" and not case.structure:
            case.structure = v

    def _normalize_factor_name(self, name: str) -> str:
        """标准化因素名称"""
        name = name.replace(' ', '').replace('\u3000', '').replace('　', '')

        mapping = {
            '繁华程度': 'prosperity',
            '楼幢位置': 'location_building',
            '朝向': 'orientation',
            '交通条件': 'traffic',
            '配套设施': 'facilities',
            '环境质量': 'environment',
            '景观': 'landscape',
            '物业管理': 'property_management',
            '驻车条件': 'parking',

            '区域位置': 'region_location',

            '建筑面积': 'area',
            '套内建筑面积': 'inner_area',
            '空间布局': 'layout',
            '新旧程度': 'age',
            '装饰装修': 'decoration',
            '建筑结构': 'structure',
            '物业类型': 'property_type',
            '设施设备': 'equipment',
            '楼宇等级': 'building_grade',
            '地形地势': 'terrain',
            '地质土壤': 'soil',
            '开发程度': 'development',

            '规划条件': 'planning',
            '土地使用期限': 'land_term',
            '土地剩余使用年限': 'land_remaining_term',
            '担保物权设立': 'mortgage',
            '租赁占用状况': 'lease',
            '拖欠税费状况': 'tax',
            '登记状况': 'registration',
            '他项权利': 'other_rights',
            '限制权利': 'restricted_rights',
            '其他因素': 'other_factors',
        }
        return mapping.get(name, name)

    # ----------------- 你 extract() 里调用到的补全方法（按 shezhi） -----------------

    def _extract_floor_factor(self, result: ZujinExtractionResult):
        """提取楼层修正系数（与 shezhi 同套路：从全文找“×xx%”）"""
        match = re.search(r'×\s*(\d+)%\s*[＝=]', self.full_text)
        if match:
            result.floor_factor = int(match.group(1)) / 100

    def _extract_extended_info(self, result: ZujinExtractionResult):
        """提取扩展信息（建成年代、价值时点、估价目的等）——按 shezhi 方式"""
        # 建成年代（租金报告不一定有，能提就提）
        build_patterns = [
            r'建成于(\d{4})年',
            r'约(\d{4})年建成',
            r'建成年代[：:]\s*(\d{4})',
            r'(\d{4})年建成',
            r'建成于上世纪(\d{2})年代',
        ]
        for pattern in build_patterns:
            match = re.search(pattern, self.full_text)
            if match:
                year_str = match.group(1)
                if len(year_str) == 2:
                    result.subject.build_year = 1900 + int(year_str)
                else:
                    result.subject.build_year = int(year_str)
                break

        # 价值时点
        value_date_patterns = [
            r'价值时点[：:]\s*(\d{4})[年\.](\d{1,2})[月\.](\d{1,2})',
            r'价值时点(\d{4})\.(\d{1,2})\.(\d{1,2})',
            r'价值时点为(\d{4})年(\d{1,2})月(\d{1,2})日',
        ]
        for pattern in value_date_patterns:
            match = re.search(pattern, self.full_text)
            if match:
                result.subject.value_date = f"{match.group(1)}-{match.group(2).zfill(2)}-{match.group(3).zfill(2)}"
                break

        # 估价目的
        purpose_patterns = [
            r'估价目的[：:是为]*(.{5,80}?)(?:。|$)',
            r'本次估价目的是(.{5,80}?)(?:。|$)',
        ]
        for pattern in purpose_patterns:
            match = re.search(pattern, self.full_text)
            if match:
                result.subject.appraisal_purpose = match.group(1).strip()
                break

    def _parse_district(self, result: ZujinExtractionResult):
        """从地址解析区域信息（照搬 shezhi 逻辑）"""
        address = result.subject.address.value or ""

        district_patterns = [
            r'([\u4e00-\u9fa5]{2,4}区)',
            r'([\u4e00-\u9fa5]{2,4}县)',
            r'([\u4e00-\u9fa5]{2,4}市)',
        ]
        for pattern in district_patterns:
            match = re.search(pattern, address)
            if match:
                result.subject.district = match.group(1)
                break

        street_patterns = [
            r'([\u4e00-\u9fa5]{2,6}街道)',
            r'([\u4e00-\u9fa5]{2,4}镇)',
            r'([\u4e00-\u9fa5]{2,4}乡)',
        ]
        for pattern in street_patterns:
            match = re.search(pattern, address)
            if match:
                result.subject.street = match.group(1)
                break

        # 同样处理可比实例
        for case in result.cases:
            case_addr = case.address.value or ""
            for pattern in district_patterns:
                match = re.search(pattern, case_addr)
                if match:
                    case.district = match.group(1)
                    break
            for pattern in street_patterns:
                match = re.search(pattern, case_addr)
                if match:
                    case.street = match.group(1)
                    break


# ============================================================================
# 测试
# ============================================================================

if __name__ == "__main__":
    extractor = ZujinExtractor()
    result = extractor.extract("./data/docs/租金报告-比较法.docx")
    print(result)

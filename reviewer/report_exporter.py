"""
审查报告导出
============
将审查结果导出为 Word 文档
"""

import os
from typing import Dict, List, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def create_review_report(review_result: Dict, output_path: str) -> str:
    """
    生成审查报告 Word 文档

    Args:
        review_result: 审查结果字典
        output_path: 输出路径

    Returns:
        生成的文件路径
    """
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)

    # 标题
    title = doc.add_heading('审查意见', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    doc.add_heading('一、基本信息', level=1)

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'

    filename = review_result.get('document_content', {}).get('filename', '未知文件')
    info_data = [
        ("报告文件", filename),
        ("审查时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("风险等级", review_result.get('overall_level', '未知')),
        ("审查摘要", review_result.get('summary', '无')),
    ]

    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)

    doc.add_paragraph()

    # 问题汇总
    doc.add_heading('二、问题汇总', level=1)

    llm_issues = review_result.get('llm_issues', [])
    validation_issues = review_result.get('validation_issues', [])
    formula_checks = review_result.get('formula_checks', [])

    # 统计
    critical_count = sum(1 for i in llm_issues if i.get('severity') == 'critical')
    major_count = sum(1 for i in llm_issues if i.get('severity') == 'major')
    minor_count = sum(1 for i in llm_issues if i.get('severity') == 'minor')
    formula_errors = sum(1 for f in formula_checks if not f.get('is_valid'))

    summary_p = doc.add_paragraph()
    summary_p.add_run(f'• 语义问题: {len(llm_issues)} 个').bold = True
    summary_p.add_run(f'（严重 {critical_count}，重要 {major_count}，轻微 {minor_count}）')

    doc.add_paragraph(f'• 校验问题: {len(validation_issues)} 个')
    doc.add_paragraph(f'• 公式异常: {formula_errors} 个')

    doc.add_paragraph()

    # 语义问题
    if llm_issues:
        doc.add_heading('三、语义问题', level=1)

        for idx, issue in enumerate(llm_issues, 1):
            severity = issue.get('severity', 'minor')
            severity_map = {'critical': '严重', 'major': '重要', 'minor': '轻微'}
            severity_cn = severity_map.get(severity, severity)

            issue_type = issue.get('type', '未知')
            paragraph_index = issue.get('paragraph_index', '')
            location = f'（段落 {paragraph_index}）' if paragraph_index else ''

            # 问题标题
            p = doc.add_paragraph()
            run = p.add_run(f'{idx}. [{severity_cn}] {issue_type} {location}')
            run.bold = True
            if severity == 'critical':
                run.font.color.rgb = RGBColor(255, 0, 0)
            elif severity == 'major':
                run.font.color.rgb = RGBColor(255, 140, 0)

            # 问题描述
            desc = issue.get('description', '')
            doc.add_paragraph(f'    问题：{desc}')

            # 原文片段
            span = issue.get('span', '')
            if span:
                span_p = doc.add_paragraph(f'   原文: ')
                span_run = span_p.add_run(f'"{span}"')
                span_run.italic = True

            # 修改建议
            suggestion = issue.get('suggestion', '')
            if suggestion:
                sug_p = doc.add_paragraph(f'   建议: ')
                sug_run = sug_p.add_run(suggestion)
                sug_run.font.color.rgb = RGBColor(0, 128, 0)

            doc.add_paragraph()

    # 校验问题
    if validation_issues:
        doc.add_heading('四、校验问题', level=1)

        val_table = doc.add_table(rows=len(validation_issues) + 1, cols=3)
        val_table.style = 'Table Grid'

        # 表头
        header_row = val_table.rows[0]
        header_row.cells[0].text = '级别'
        header_row.cells[1].text = '类别'
        header_row.cells[2].text = '描述'

        for cell in header_row.cells:
            cell.paragraphs[0].runs[0].bold = True

        # 数据
        for i, issue in enumerate(validation_issues, 1):
            row = val_table.rows[i]
            row.cells[0].text = issue.get('level', '')
            row.cells[1].text = issue.get('category', '')
            row.cells[2].text = issue.get('description', '')

        doc.add_paragraph()

    # 公式校验
    if formula_checks:
        doc.add_heading('五、公式校验', level=1)

        formula_table = doc.add_table(rows=len(formula_checks) + 1, cols=4)
        formula_table.style = 'Table Grid'

        # 表头
        header_row = formula_table.rows[0]
        headers = ['案例', '预期值', '实际值', '结果']
        for j, h in enumerate(headers):
            header_row.cells[j].text = h
            header_row.cells[j].paragraphs[0].runs[0].bold = True

        # 数据
        for i, fc in enumerate(formula_checks, 1):
            row = formula_table.rows[i]
            row.cells[0].text = fc.get('case_id', '')
            row.cells[1].text = f"{fc.get('expected', 0):,.2f}"
            row.cells[2].text = f"{fc.get('actual', 0):,.2f}"
            result = '通过' if fc.get('is_valid') else '异常'
            row.cells[3].text = result
            if not fc.get('is_valid'):
                row.cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)

        doc.add_paragraph()

    # 审查结论
    doc.add_heading('六、审查结论', level=1)

    risk = review_result.get('overall_level', '未知')
    if risk == '高风险':
        conclusion = '该报告存在较多严重问题，建议退回修改后重新提交。'
    elif risk == '中风险':
        conclusion = '该报告存在一些问题，建议修改完善后再使用。'
    else:
        conclusion = '该报告整体质量较好，可以使用。'

    doc.add_paragraph(conclusion)

    doc.add_paragraph()
    doc.add_paragraph()

    # 签名区
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_p.add_run('审核人：_______________________')

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run(f'日期：{datetime.now().strftime("%Y年%m月%d日")}')

    # 保存
    doc.save(output_path)
    return output_path


def create_review_report_with_original(review_result: Dict, output_path: str) -> str:
    """
    生成带原文标注的审查报告

    Args:
        review_result: 审查结果字典（包含 document_content）
        output_path: 输出路径

    Returns:
        生成的文件路径
    """
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)

    # 标题
    title = doc.add_heading('房地产估价报告审查意见（含原文标注）', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    doc.add_heading('一、基本信息', level=1)

    filename = review_result.get('document_content', {}).get('filename', '未知文件')
    doc.add_paragraph(f'报告文件：{filename}')
    doc.add_paragraph(f'审查时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'风险等级：{review_result.get("overall_risk", "未知")}')
    doc.add_paragraph(f'审查摘要：{review_result.get("summary", "")}')

    doc.add_paragraph()

    # 问题汇总（简版）
    doc.add_heading('二、问题汇总', level=1)

    llm_issues = review_result.get('llm_issues', [])
    doc.add_paragraph(f'共发现 {len(llm_issues)} 个语义问题，详见原文标注。')

    doc.add_paragraph()

    # 原文内容（带标注）
    doc.add_heading('三、原文内容（问题段落已标注）', level=1)

    document_content = review_result.get('document_content', {})
    contents = document_content.get('contents', [])

    # 构建段落索引到问题的映射
    issue_map = {}
    for issue in llm_issues:
        p_idx = issue.get('paragraph_index')
        if p_idx is not None:
            if p_idx not in issue_map:
                issue_map[p_idx] = []
            issue_map[p_idx].append(issue)

    for item in contents:
        if item.get('type') == 'paragraph':
            idx = item.get('index')
            text = item.get('text', '')
            has_issue = item.get('has_issue', False)

            p = doc.add_paragraph()

            # 段落编号
            index_run = p.add_run(f'[{idx}] ')
            index_run.font.color.rgb = RGBColor(128, 128, 128)
            index_run.font.size = Pt(10)

            # 段落内容
            text_run = p.add_run(text)

            if has_issue:
                # 高亮问题段落
                text_run.font.color.rgb = RGBColor(255, 0, 0)

                # 添加问题批注
                issues = issue_map.get(idx, [])
                for issue in issues:
                    comment_p = doc.add_paragraph()
                    comment_p.paragraph_format.left_indent = Inches(0.5)

                    severity = issue.get('severity', 'minor')
                    severity_map = {'critical': '严重', 'major': '重要', 'minor': '轻微'}

                    comment_run = comment_p.add_run(
                        f'⚠ [{severity_map.get(severity, severity)}] {issue.get("description", "")}'
                    )
                    comment_run.font.size = Pt(10)
                    comment_run.font.color.rgb = RGBColor(255, 140, 0)

                    if issue.get('suggestion'):
                        sug_p = doc.add_paragraph()
                        sug_p.paragraph_format.left_indent = Inches(0.5)
                        sug_run = sug_p.add_run(f'  建议：{issue.get("suggestion")}')
                        sug_run.font.size = Pt(10)
                        sug_run.font.color.rgb = RGBColor(0, 128, 0)

        elif item.get('type') == 'table':
            # 表格简化处理
            doc.add_paragraph('[表格内容略]')

    # 保存
    doc.save(output_path)
    return output_path
